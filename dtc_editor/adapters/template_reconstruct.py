"""
Template-based document reconstruction for DTC compliance.

This module provides functionality to:
1. Detect if a document follows the DTC template
2. Extract content from non-compliant documents
3. Reconstruct documents using the official DTC template
"""

from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
import re
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# DTC Template style names
DTC_STYLES = {
    "title": "Title",
    "heading1": "Heading 1",
    "heading2": "Heading 2",
    "heading3": "Heading 3",
    "normal": "Normal",
    "bullet": "List Bullet 2",
    "caption": "Caption",
    "toc_heading": "TOC Heading",
    "table_of_figures": "table of figures",
}

# Required DTC styles for compliance detection
REQUIRED_DTC_STYLES = {"Heading 1", "Heading 2", "Caption", "List Bullet 2"}


@dataclass
class ExtractedFigure:
    """A figure extracted from the source document."""
    index: int
    drawing_element: Any  # The XML element
    caption: Optional[str] = None
    section_number: str = "1"
    figure_number: int = 1


@dataclass
class ExtractedTable:
    """A table extracted from the source document."""
    index: int
    table: Any  # The docx Table object
    caption: Optional[str] = None
    section_number: str = "1"
    table_number: int = 1


@dataclass
class ExtractedHeading:
    """A heading extracted from the source document."""
    text: str
    level: int  # 1, 2, or 3
    index: int


@dataclass
class ExtractedParagraph:
    """A paragraph extracted from the source document."""
    text: str
    style_hint: str  # "normal", "bullet", "abstract"
    index: int
    runs: List[Tuple[str, Dict]] = field(default_factory=list)  # (text, formatting)


@dataclass
class ExtractedContent:
    """All content extracted from a source document."""
    title: str = ""
    date: str = ""
    authors: List[str] = field(default_factory=list)
    abstract: str = ""
    headings: List[ExtractedHeading] = field(default_factory=list)
    paragraphs: List[ExtractedParagraph] = field(default_factory=list)
    figures: List[ExtractedFigure] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    references: List[str] = field(default_factory=list)

    # Content ordering for reconstruction
    content_order: List[Tuple[str, int]] = field(default_factory=list)  # ("heading", idx), ("para", idx), etc.


@dataclass
class TemplateComplianceResult:
    """Result of template compliance check."""
    is_compliant: bool
    score: float  # 0.0 to 1.0
    missing_styles: List[str] = field(default_factory=list)
    issues: List[str] = field(default_factory=list)
    recommendation: str = ""  # "full_reconstruct" | "light_fix" | "compliant"


def detect_template_compliance(doc: Document) -> TemplateComplianceResult:
    """
    Check if a document follows the DTC template.

    Returns a compliance result with score and recommendations.
    """
    issues = []
    missing_styles = []

    # Get all style names used in the document
    used_styles = set()
    for para in doc.paragraphs:
        if para.style:
            used_styles.add(para.style.name)

    # Check for required DTC styles
    for required_style in REQUIRED_DTC_STYLES:
        if required_style not in used_styles:
            missing_styles.append(required_style)

    # Check for DTC-specific styles in document's style definitions
    doc_style_names = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH}
    has_dtc_styles = any("DTC" in name for name in doc_style_names)

    # Check title page elements
    has_title = any(p.style and p.style.name == "Title" for p in doc.paragraphs[:10])

    # Check for proper heading hierarchy
    heading_levels_used = set()
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
                heading_levels_used.add(level)
            except ValueError:
                pass

    has_heading_hierarchy = 1 in heading_levels_used

    # Check for caption style usage
    has_captions = "Caption" in used_styles or "DTC figure caption" in used_styles

    # Calculate compliance score
    score = 0.0
    checks = [
        (len(missing_styles) == 0, 0.3, "Missing required styles"),
        (has_dtc_styles, 0.2, "No DTC-specific styles found"),
        (has_title, 0.15, "No Title style on title page"),
        (has_heading_hierarchy, 0.2, "No Heading 1 style used"),
        (has_captions, 0.15, "No Caption style used"),
    ]

    for passed, weight, issue_msg in checks:
        if passed:
            score += weight
        else:
            issues.append(issue_msg)

    # Determine recommendation
    if score >= 0.8:
        recommendation = "compliant"
        is_compliant = True
    elif score >= 0.5:
        recommendation = "light_fix"
        is_compliant = False
    else:
        recommendation = "full_reconstruct"
        is_compliant = False

    return TemplateComplianceResult(
        is_compliant=is_compliant,
        score=score,
        missing_styles=missing_styles,
        issues=issues,
        recommendation=recommendation,
    )


def _get_paragraph_font_size(para) -> Optional[float]:
    """Get the effective font size of a paragraph."""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    # Check style
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size.pt
    return None


def _is_paragraph_bold(para) -> bool:
    """Check if paragraph text is bold."""
    for run in para.runs:
        if run.bold or (run.font and run.font.bold):
            return True
    return False


def _infer_heading_level(para, body_font_size: float = 12.0) -> Optional[int]:
    """Infer heading level from formatting."""
    font_size = _get_paragraph_font_size(para)
    is_bold = _is_paragraph_bold(para)
    text = para.text.strip()
    word_count = len(text.split())

    if not text or word_count > 15:
        return None

    # Check for numbered heading pattern
    numbered_pattern = re.match(r'^(\d+(?:\.\d+)*)\s+', text)

    # Score-based detection
    score = 0
    if font_size and font_size > body_font_size:
        score += 3
    if is_bold:
        score += 2
    if numbered_pattern:
        score += 2
    if word_count <= 7:
        score += 1

    if score < 3:
        return None

    # Determine level
    if numbered_pattern:
        num_parts = numbered_pattern.group(1).count('.') + 1
        return min(num_parts, 3)

    if font_size:
        if font_size >= 16:
            return 1
        elif font_size >= 14:
            return 2
        else:
            return 3

    return 1 if is_bold else None


def _is_bullet_paragraph(para) -> bool:
    """Check if paragraph is a bullet point."""
    style_name = para.style.name if para.style else ""
    if "bullet" in style_name.lower() or "list" in style_name.lower():
        return True

    # Check for bullet characters
    text = para.text.strip()
    if text and text[0] in "•●○■□▪▫-–—":
        return True

    # Check XML for numbering
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            return True

    return False


def _extract_run_formatting(run) -> Dict:
    """Extract formatting from a run."""
    return {
        "bold": run.bold or False,
        "italic": run.italic or False,
        "underline": run.underline or False,
        "font_name": run.font.name,
        "font_size": run.font.size.pt if run.font.size else None,
    }


def extract_document_content(doc: Document) -> ExtractedContent:
    """
    Extract all content from a source document.

    This extracts text, figures, tables, and maintains ordering
    for reconstruction.
    """
    content = ExtractedContent()

    # Determine body font size (most common)
    font_sizes = []
    for para in doc.paragraphs:
        size = _get_paragraph_font_size(para)
        if size:
            font_sizes.append(size)
    body_font_size = max(set(font_sizes), key=font_sizes.count) if font_sizes else 12.0

    # Track current section for figure/table numbering
    current_section = 1
    figure_counter = 0
    table_counter = 0

    # Track if we've passed the first real content heading (for abstract detection)
    first_heading_seen = False
    abstract_paragraphs = []
    in_abstract_section = False

    # Track table indices to avoid double-processing
    processed_tables = set()

    # Front matter patterns to skip
    front_matter_patterns = [
        r'^white\s*paper$',
        r'^tech(nical)?\s*brief$',
        r'^best\s*practices?\s*paper$',
        r'^date[:\s]',
        r'^authors?[:\s]?$',
        r'^contributors?[:\s]?$',
    ]

    # Process paragraphs
    para_idx = 0
    heading_idx = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        # Skip front matter patterns
        is_front_matter = any(re.match(p, text.lower()) for p in front_matter_patterns)

        # Check for title (first paragraphs with large font, not front matter labels)
        if i < 20 and not is_front_matter:
            style_name = para.style.name if para.style else ""
            font_size = _get_paragraph_font_size(para)
            # Title is usually larger than body text (14pt+), or has Title style
            if style_name == "Title" or (font_size and font_size >= 14 and font_size > body_font_size):
                # Skip generic labels
                if text.lower() not in ["white paper", "tech brief", "technical brief"]:
                    # Handle subtitle in parentheses (e.g., "("MEC")")
                    if text.startswith("(") and text.endswith(")") and content.title:
                        # Append as subtitle
                        content.title = f"{content.title} {text}"
                        continue
                    elif not content.title:
                        content.title = text
                        continue

        # Check for date pattern
        if i < 25 and not content.date:
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})|(\d{1,2}/\d{1,2}/\d{4})|((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})', text, re.IGNORECASE)
            if date_match:
                content.date = date_match.group(0)
                continue

        # Check for authors line (look for "Name (Organization)" pattern)
        if i < 30 and not content.authors and not is_front_matter:
            # Pattern: Name (Org), Name (Org)
            author_pattern = re.findall(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*\([^)]+\)', text)
            if author_pattern and len(author_pattern) >= 1:
                content.authors = author_pattern
                continue

        # Check for heading
        heading_level = _infer_heading_level(para, body_font_size)
        if heading_level:
            # Skip front matter headings and subtitles
            text_lower = text.lower()
            if text_lower in ["white paper", "tech brief", "technical brief", "best practices paper"]:
                continue
            # Skip parenthetical subtitles like ("MEC")
            if text.startswith("(") and text.endswith(")"):
                continue

            # Check if this is the Abstract heading
            if text_lower == "abstract":
                in_abstract_section = True
                first_heading_seen = True
                continue  # Don't include "Abstract" as a heading, we'll add it in reconstruction

            # Any heading after Abstract ends the abstract section
            if in_abstract_section:
                in_abstract_section = False

            first_heading_seen = True
            if heading_level == 1:
                current_section += 1
                figure_counter = 0
                table_counter = 0

            heading = ExtractedHeading(
                text=text,
                level=heading_level,
                index=heading_idx,
            )
            content.headings.append(heading)
            content.content_order.append(("heading", heading_idx))
            heading_idx += 1
            continue

        # Abstract detection: paragraphs in abstract section OR before first content heading
        if in_abstract_section or (not first_heading_seen and len(text) > 30):
            abstract_paragraphs.append(text)
            continue

        # Check for reference entries
        if re.match(r'^\[\d+\]', text):
            content.references.append(text)
            continue

        # Regular paragraph or bullet
        style_hint = "bullet" if _is_bullet_paragraph(para) else "normal"

        # Extract runs with formatting
        runs = []
        for run in para.runs:
            if run.text:
                runs.append((run.text, _extract_run_formatting(run)))

        paragraph = ExtractedParagraph(
            text=text,
            style_hint=style_hint,
            index=para_idx,
            runs=runs,
        )
        content.paragraphs.append(paragraph)
        content.content_order.append(("para", para_idx))
        para_idx += 1

    # Set abstract
    if abstract_paragraphs:
        content.abstract = " ".join(abstract_paragraphs)

    # Extract figures (drawings)
    figure_idx = 0
    current_section = 1
    for i, para in enumerate(doc.paragraphs):
        # Check for section change
        if para.style and para.style.name == "Heading 1":
            current_section += 1

        # Find drawings in paragraph
        drawings = para._element.findall('.//' + qn('w:drawing'))
        for drawing in drawings:
            figure_counter += 1

            # Try to find caption in next paragraph
            caption = None
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if next_text.lower().startswith("figure"):
                    caption = next_text

            figure = ExtractedFigure(
                index=figure_idx,
                drawing_element=drawing,
                caption=caption,
                section_number=str(current_section),
                figure_number=figure_counter,
            )
            content.figures.append(figure)
            figure_idx += 1

    # Extract tables
    table_idx = 0
    current_section = 1
    table_counter = 0

    for table in doc.tables:
        table_counter += 1

        extracted_table = ExtractedTable(
            index=table_idx,
            table=table,
            caption=None,  # Will need to find from context
            section_number=str(current_section),
            table_number=table_counter,
        )
        content.tables.append(extracted_table)
        table_idx += 1

    return content


def _copy_table(source_table, target_doc) -> Any:
    """Copy a table from source to target document."""
    # Create new table with same dimensions
    rows = len(source_table.rows)
    cols = len(source_table.columns)

    new_table = target_doc.add_table(rows=rows, cols=cols)

    # Copy cell contents
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_table.rows[i].cells[j].text = cell.text

    # Try to copy table style
    if source_table.style:
        try:
            new_table.style = source_table.style
        except:
            pass

    return new_table


def _copy_drawing_to_paragraph(drawing_element, target_para):
    """Copy a drawing element to a target paragraph."""
    # Create a new run in the target paragraph
    run = target_para.add_run()

    # Deep copy the drawing element
    new_drawing = copy.deepcopy(drawing_element)

    # Add to run's XML
    run._element.append(new_drawing)


def reconstruct_from_template(
    content: ExtractedContent,
    template_path: str,
    output_path: str,
    llm_client: Optional[Any] = None,
) -> Dict[str, Any]:
    """
    Reconstruct a document using the DTC template.

    Args:
        content: Extracted content from source document
        template_path: Path to DTC template file
        output_path: Path for output document
        llm_client: Optional LLM client for caption inference

    Returns:
        Dictionary with reconstruction statistics
    """
    # Load template
    doc = Document(template_path)

    stats = {
        "headings_added": 0,
        "paragraphs_added": 0,
        "figures_added": 0,
        "tables_added": 0,
        "references_added": 0,
    }

    # Find key template elements to preserve
    title_para = None
    date_para = None
    authors_para = None
    tof_heading = None
    tot_heading = None

    # Identify template structure elements
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if style_name == "Title" and not title_para:
            title_para = para
        elif "YYYY-MM-DD" in text and not date_para:
            date_para = para
        elif "Name (Organization)" in text and not authors_para:
            authors_para = para
        elif style_name == "TOC Heading" and text == "Figures":
            tof_heading = para
        elif style_name == "TOC Heading" and text == "Tables":
            tot_heading = para

    # Remove ALL paragraphs except essential front matter
    # Keep: title, date placeholder, authors placeholder, TOC headings
    elements_to_keep = {id(p._element) for p in [title_para, date_para, authors_para, tof_heading, tot_heading] if p}

    paragraphs_to_remove = []
    for para in doc.paragraphs:
        if id(para._element) not in elements_to_keep:
            # Check if it's an essential structural element
            style_name = para.style.name if para.style else ""
            # Keep section breaks and empty spacing paragraphs in front matter
            if style_name in ["TOC Heading"]:
                continue
            paragraphs_to_remove.append(para)

    # Remove all template content paragraphs
    for para in paragraphs_to_remove:
        p = para._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    # Remove all template tables
    for table in doc.tables:
        tbl = table._element
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)

    # Update title page
    if title_para:
        title_para.clear()
        title_para.add_run(content.title or "Untitled Document")

    # Update date
    if date_para:
        date_para.clear()
        date_para.add_run(content.date or "2025-01-14")

    # Update authors
    if authors_para:
        authors_para.clear()
        authors_text = ", ".join(content.authors) if content.authors else "Author (Organization)"
        authors_para.add_run(authors_text)

    # Add abstract if present
    if content.abstract:
        abstract_para = doc.add_paragraph(content.abstract)
        try:
            abstract_para.style = doc.styles["Normal"]
        except KeyError:
            pass
        stats["paragraphs_added"] += 1

    # Track current section for figure/table numbering
    current_section = 0
    figure_counter = 0
    table_counter = 0

    # Process content in order
    heading_map = {h.index: h for h in content.headings}
    para_map = {p.index: p for p in content.paragraphs}

    for content_type, idx in content.content_order:
        if content_type == "heading":
            heading = heading_map.get(idx)
            if heading:
                if heading.level == 1:
                    current_section += 1
                    figure_counter = 0
                    table_counter = 0

                style_name = f"Heading {heading.level}"
                para = doc.add_paragraph(heading.text)
                try:
                    para.style = doc.styles[style_name]
                except KeyError:
                    para.style = doc.styles["Heading 1"]
                stats["headings_added"] += 1

        elif content_type == "para":
            paragraph = para_map.get(idx)
            if paragraph:
                para = doc.add_paragraph()

                # Apply style
                style_name = "List Bullet 2" if paragraph.style_hint == "bullet" else "Normal"
                try:
                    para.style = doc.styles[style_name]
                except KeyError:
                    para.style = doc.styles["Normal"]

                # Add text with formatting
                if paragraph.runs:
                    for text, fmt in paragraph.runs:
                        run = para.add_run(text)
                        if fmt.get("bold"):
                            run.bold = True
                        if fmt.get("italic"):
                            run.italic = True
                else:
                    para.add_run(paragraph.text)

                stats["paragraphs_added"] += 1

    # Add figures with captions
    for figure in content.figures:
        # Add figure
        fig_para = doc.add_paragraph()
        fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _copy_drawing_to_paragraph(figure.drawing_element, fig_para)

        # Add caption
        caption_para = doc.add_paragraph()
        try:
            caption_para.style = doc.styles["Caption"]
        except KeyError:
            pass

        figure_counter += 1
        section_num = current_section if current_section > 0 else 1

        if figure.caption:
            caption_text = figure.caption
        else:
            caption_text = f"Figure {section_num}-{figure_counter}: SOURCE AND/OR DESCRIPTION NEEDED."
            # Make placeholder red
            run = caption_para.add_run(caption_text)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True
            stats["figures_added"] += 1
            continue

        caption_para.add_run(caption_text)
        stats["figures_added"] += 1

    # Add tables with captions
    for table_info in content.tables:
        table_counter += 1
        section_num = current_section if current_section > 0 else 1

        # Add caption before table
        caption_para = doc.add_paragraph()
        try:
            caption_para.style = doc.styles["Caption"]
        except KeyError:
            pass

        if table_info.caption:
            caption_para.add_run(table_info.caption)
        else:
            caption_text = f"Table {section_num}-{table_counter}: Table description."
            caption_para.add_run(caption_text)

        # Copy table
        _copy_table(table_info.table, doc)
        stats["tables_added"] += 1

    # Add References section
    if content.references:
        ref_heading = doc.add_paragraph("References")
        try:
            ref_heading.style = doc.styles["Heading 1"]
        except KeyError:
            pass

        for ref in content.references:
            ref_para = doc.add_paragraph(ref)
            try:
                ref_para.style = doc.styles["Normal"]
            except KeyError:
                pass
            stats["references_added"] += 1

    # Save document
    doc.save(output_path)

    return stats


def reconstruct_document(
    source_path: str,
    template_path: str,
    output_path: str,
    force_reconstruct: bool = False,
    llm_client: Optional[Any] = None,
) -> Dict[str, Any]:
    """
    Main entry point for template-based reconstruction.

    Args:
        source_path: Path to source document
        template_path: Path to DTC template
        output_path: Path for output document
        force_reconstruct: If True, always do full reconstruction
        llm_client: Optional LLM client for caption inference

    Returns:
        Dictionary with reconstruction results and statistics
    """
    # Load source document
    source_doc = Document(source_path)

    # Check compliance
    compliance = detect_template_compliance(source_doc)

    result = {
        "source_path": source_path,
        "output_path": output_path,
        "compliance_score": compliance.score,
        "compliance_issues": compliance.issues,
        "recommendation": compliance.recommendation,
        "action_taken": "",
        "stats": {},
    }

    if compliance.is_compliant and not force_reconstruct:
        # Document is compliant, just copy it
        source_doc.save(output_path)
        result["action_taken"] = "copied"
        result["stats"] = {"note": "Document already compliant"}
        return result

    if compliance.recommendation == "light_fix" and not force_reconstruct:
        # TODO: Implement light fixes (style remapping)
        # For now, fall through to full reconstruction
        pass

    # Full reconstruction
    content = extract_document_content(source_doc)
    stats = reconstruct_from_template(
        content=content,
        template_path=template_path,
        output_path=output_path,
        llm_client=llm_client,
    )

    result["action_taken"] = "reconstructed"
    result["stats"] = stats
    result["extracted"] = {
        "title": content.title,
        "authors": content.authors,
        "headings": len(content.headings),
        "paragraphs": len(content.paragraphs),
        "figures": len(content.figures),
        "tables": len(content.tables),
        "references": len(content.references),
    }

    return result
