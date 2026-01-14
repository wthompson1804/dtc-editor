from __future__ import annotations
from typing import Set, Tuple, List
from docx import Document
from dtc_editor.ir import DocumentIR, TextBlock, BlockRef, StructureInventory
import hashlib

def _norm(s: str) -> str:
    return " ".join(s.split()).strip()

def _anchor(prev_txt: str, txt: str, next_txt: str) -> str:
    # Stable-ish: hash of neighbors + normalized text
    h = hashlib.sha256()
    h.update((_norm(prev_txt)[:120] + "\n" + _norm(txt)[:500] + "\n" + _norm(next_txt)[:120]).encode("utf-8"))
    return h.hexdigest()[:16]

def extract_ir_and_inventory(docx_path: str) -> Tuple[DocumentIR, StructureInventory]:
    doc = Document(docx_path)
    ir = DocumentIR()
    inv = StructureInventory()
    inv.table_count = len(doc.tables)

    # title
    for p in doc.paragraphs:
        if p.text.strip():
            if p.style and p.style.name == "Title":
                ir.title = p.text.strip()
                break
    if not ir.title:
        for p in doc.paragraphs:
            if p.text.strip():
                ir.title = p.text.strip()
                break

    # Build blocks with anchors
    texts = [p.text for p in doc.paragraphs]
    doc_i = 0
    p_i = 0
    h_i = 0
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text
        style = p.style.name if p.style else ""
        if not txt:
            doc_i += 1
            continue

        is_heading = style.lower().startswith("heading")
        bt = "heading" if is_heading else "paragraph"
        bi = h_i if is_heading else p_i
        prev_txt = texts[idx-1] if idx-1 >= 0 else ""
        next_txt = texts[idx+1] if idx+1 < len(texts) else ""
        anch = _anchor(prev_txt, txt, next_txt)

        ir.blocks.append(TextBlock(
            ref=BlockRef(block_type=bt, doc_index=idx, block_index=bi),
            style_name=style,
            text=txt,
            anchor=anch,
        ))

        if is_heading:
            inv.headings.append(txt.strip())
            inv.heading_styles.append(style)
            h_i += 1
        else:
            p_i += 1
        doc_i += 1

    inv.paragraph_count = len([b for b in ir.blocks if b.ref.block_type == "paragraph"])
    whole = "\n\n".join([b.text for b in ir.blocks]).lower()
    inv.has_abstract = "abstract" in whole
    inv.has_references = "references" in whole
    inv.has_authors = "authors" in whole
    return ir, inv

def emit_clean_docx(original_docx: str, updated_ir: DocumentIR, out_docx: str) -> None:
    doc = Document(original_docx)
    # apply by doc_index (fallback) - best effort
    for block in updated_ir.blocks:
        p = doc.paragraphs[block.ref.doc_index]
        if p.text != block.text:
            p.text = block.text
    doc.save(out_docx)

def load_protected_terms(path: str) -> Set[str]:
    import yaml
    with open(path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    return set([t.strip() for t in (data.get("protected_terms") or []) if isinstance(t,str) and t.strip()])
