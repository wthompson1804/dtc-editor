"""
Document Restructuring Module for DTC Editorial Engine.

Transforms non-compliant documents to follow DTC template structure:
- Adds Table of Contents
- Adds Table of Figures
- Adds Table of Tables
- Inserts required DTC sections
- Applies proper heading hierarchy
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy


@dataclass
class DocumentSection:
    """Represents a section of the document."""
    title: str
    level: int  # 1 = Heading 1, 2 = Heading 2, etc.
    content_paragraphs: List[int] = field(default_factory=list)  # indices
    subsections: List["DocumentSection"] = field(default_factory=list)


@dataclass
class DocumentAnalysis:
    """Analysis of document structure."""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    abstract: Optional[str] = None
    sections: List[DocumentSection] = field(default_factory=list)
    figure_count: int = 0
    table_count: int = 0
    has_toc: bool = False
    has_tof: bool = False
    has_tot: bool = False
    has_references: bool = False
    missing_elements: List[str] = field(default_factory=list)


@dataclass
class RestructureConfig:
    """Configuration for document restructuring."""
    template_path: Optional[str] = None
    add_toc: bool = True
    add_tof: bool = True
    add_tot: bool = True
    add_title_page: bool = True
    document_type: str = "White Paper"  # White Paper, Tech Brief, Best Practice
    date_string: Optional[str] = None
    use_inferred_content: bool = True  # True = actual text, False = Word fields
    include_missing_placeholders: bool = True  # Add red placeholders for missing sections


@dataclass
class TOCEntry:
    """Entry in table of contents."""
    text: str
    level: int
    page_hint: str = ""  # Can't know actual page, but can use section number


@dataclass
class CaptionEntry:
    """Entry for figure or table caption."""
    label: str      # "Figure 1-1" or "Table 2-1"
    caption: str    # The caption text
    page_hint: str = ""


@dataclass
class InferredStructure:
    """Inferred document structure for TOC/TOF/TOT generation."""
    toc_entries: List[TOCEntry] = field(default_factory=list)
    figure_entries: List[CaptionEntry] = field(default_factory=list)
    table_entries: List[CaptionEntry] = field(default_factory=list)


def _get_paragraph_font_size(p) -> Optional[float]:
    """Get the font size of a paragraph (from first run or style)."""
    # Try to get from runs
    for run in p.runs:
        if run.font.size:
            return run.font.size.pt
    # Try to get from style
    if p.style and p.style.font and p.style.font.size:
        return p.style.font.size.pt
    return None


def _is_paragraph_bold(p) -> bool:
    """Check if paragraph is bold (from runs or style)."""
    # Check runs
    for run in p.runs:
        if run.bold or (run.font and run.font.bold):
            return True
    # Check style
    if p.style and p.style.font and p.style.font.bold:
        return True
    return False


def _get_body_font_size(doc: Document) -> float:
    """Estimate the body text font size by finding the most common size."""
    from collections import Counter
    sizes = Counter()

    for p in doc.paragraphs:
        if len(p.text.strip()) > 50:  # Likely body text
            size = _get_paragraph_font_size(p)
            if size:
                sizes[size] += 1

    if sizes:
        return sizes.most_common(1)[0][0]
    return 12.0  # Default assumption


def infer_document_structure(doc: Document) -> InferredStructure:
    """
    Analyze document to infer TOC, TOF, and TOT entries.

    Uses formatting-based detection:
    - Larger font size than body text
    - Bold text
    - Short length (< 10 words)
    - No ending punctuation
    - Preceded/followed by spacing

    Also detects figure/table captions for TOF/TOT.
    """
    structure = InferredStructure()
    current_chapter = 1

    # Determine body font size for comparison
    body_font_size = _get_body_font_size(doc)

    # First pass: check if document uses Heading styles
    has_styled_headings = any(
        p.style and p.style.name.lower().startswith("heading")
        for p in doc.paragraphs
    )

    paragraphs = doc.paragraphs
    num_paragraphs = len(paragraphs)

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""

        if not text:
            continue

        # Skip front matter labels we're generating
        if text.lower() in ["contents", "figures", "tables", "table of contents",
                            "table of figures", "table of tables"]:
            continue

        # FIRST: Detect figure/table captions (before any continue statements)
        fig_match = re.match(r'^(Figure\s+\d+(?:-\d+)?)[:\.]?\s*(.*)$', text, re.IGNORECASE)
        if fig_match:
            structure.figure_entries.append(CaptionEntry(
                label=fig_match.group(1),
                caption=fig_match.group(2).rstrip('.'),
            ))
            continue  # Don't also treat as heading

        tbl_match = re.match(r'^(Table\s+\d+(?:-\d+)?)[:\.]?\s*(.*)$', text, re.IGNORECASE)
        if tbl_match:
            structure.table_entries.append(CaptionEntry(
                label=tbl_match.group(1),
                caption=tbl_match.group(2).rstrip('.'),
            ))
            continue  # Don't also treat as heading

        # THEN: Detect headings for TOC
        is_heading = False
        level = 1

        # Method 1: From Heading styles (most reliable)
        if style.lower().startswith("heading"):
            is_heading = True
            if "2" in style:
                level = 2
            elif "3" in style:
                level = 3

        # Method 2: From formatting characteristics
        elif not has_styled_headings:
            words = text.split()
            word_count = len(words)

            # Skip list items
            if "list" in style.lower():
                continue

            # Basic requirements for a heading
            is_short = word_count <= 10
            no_end_punct = not text.endswith('.') and not text.endswith(',') and not text.endswith(':')

            if not (is_short and no_end_punct):
                continue

            # Skip obvious non-headings
            skip_patterns = [
                r'^\(.*\)$',           # Parenthetical like "(MEC)"
                r'^Date:',             # Metadata
                r'^Authors?:',         # Metadata
                r'^Version:',          # Metadata
                r'^\d{4}[-/]',         # Date patterns
                r'^(http|www\.)',      # URLs
            ]
            if any(re.match(pat, text, re.IGNORECASE) for pat in skip_patterns):
                continue

            # Get formatting info
            font_size = _get_paragraph_font_size(p)
            is_bold = _is_paragraph_bold(p)

            # Check for visual separation (blank line before or after)
            prev_empty = i == 0 or not paragraphs[i-1].text.strip()
            next_empty = i >= num_paragraphs - 1 or not paragraphs[i+1].text.strip()
            has_spacing = prev_empty or next_empty

            # Heading detection heuristics
            is_larger = font_size and font_size > body_font_size + 0.5
            is_numbered = bool(re.match(r'^(\d+(?:\.\d+)*)\s+\S', text))

            # Score-based detection
            heading_score = 0
            if is_larger:
                heading_score += 3
            if is_bold:
                heading_score += 2
            if has_spacing:
                heading_score += 1
            if is_numbered:
                heading_score += 2
            if word_count <= 5:
                heading_score += 1

            # Threshold: needs strong evidence
            if heading_score >= 3:
                is_heading = True

                # Determine level from font size or numbering
                if is_numbered:
                    parts = re.match(r'^(\d+(?:\.\d+)*)', text).group(1).split('.')
                    level = min(len(parts), 3)
                elif font_size:
                    # Larger = higher level (lower number)
                    if font_size >= body_font_size + 4:
                        level = 1
                    elif font_size >= body_font_size + 2:
                        level = 2
                    else:
                        level = 3

        if is_heading and text:
            # Extract section number if present
            match = re.match(r'^(\d+(?:\.\d+)*)\s*', text)
            if match:
                page_hint = match.group(1)
            elif level == 1:
                page_hint = str(current_chapter)
                current_chapter += 1
            else:
                page_hint = ""

            structure.toc_entries.append(TOCEntry(
                text=text,
                level=level,
                page_hint=page_hint,
            ))

    return structure


def analyze_document(doc: Document) -> DocumentAnalysis:
    """
    Analyze document structure to identify what's present and missing.
    """
    analysis = DocumentAnalysis()

    # Find title
    for p in doc.paragraphs[:10]:
        if p.style and p.style.name == "Title" and p.text.strip():
            analysis.title = p.text.strip()
            break
        elif p.text.strip() and not analysis.title:
            analysis.title = p.text.strip()

    # Scan for structure
    current_section = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        style = p.style.name if p.style else ""

        # Check for TOC
        if "table of contents" in text or "toc heading" in style.lower():
            analysis.has_toc = True

        # Check for TOF
        if "table of figures" in text or "figures" == text:
            analysis.has_tof = True

        # Check for TOT
        if "table of tables" in text or "tables" == text:
            analysis.has_tot = True

        # Check for Abstract
        if text in ["abstract", "executive summary"]:
            analysis.abstract = ""
            # Collect abstract text from following paragraphs
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                next_p = doc.paragraphs[j]
                if next_p.style and "heading" in next_p.style.name.lower():
                    break
                if next_p.text.strip():
                    analysis.abstract = (analysis.abstract or "") + next_p.text.strip() + " "

        # Check for References
        if text in ["references", "bibliography"]:
            analysis.has_references = True

        # Build section structure
        if style.lower().startswith("heading"):
            level = 1
            if "2" in style:
                level = 2
            elif "3" in style:
                level = 3

            section = DocumentSection(
                title=p.text.strip(),
                level=level,
            )
            analysis.sections.append(section)

    # Count figures and tables
    for p in doc.paragraphs:
        drawings = p._element.findall('.//' + qn('w:drawing'))
        analysis.figure_count += len(drawings)

    analysis.table_count = len(doc.tables)

    # Identify missing elements
    if not analysis.has_toc:
        analysis.missing_elements.append("Table of Contents")
    if not analysis.has_tof and analysis.figure_count > 0:
        analysis.missing_elements.append("Table of Figures")
    if not analysis.has_tot and analysis.table_count > 0:
        analysis.missing_elements.append("Table of Tables")
    if not analysis.abstract:
        analysis.missing_elements.append("Abstract/Executive Summary")
    if not analysis.has_references:
        analysis.missing_elements.append("References section")

    return analysis


def insert_inferred_toc(
    doc: Document,
    entries: List[TOCEntry],
    insert_position: int,
) -> int:
    """
    Insert an inferred Table of Contents as actual text.

    Returns number of paragraphs inserted.
    """
    if not entries:
        return 0

    body = doc.element.body
    inserted = 0

    # Add heading
    heading = doc.add_paragraph()
    heading.style = "Heading 1"
    heading.add_run("Contents")
    body.insert(insert_position, heading._element)
    inserted += 1

    # Add entries
    for entry in entries:
        p = doc.add_paragraph()

        # Indent based on level
        indent = "    " * (entry.level - 1)
        text = f"{indent}{entry.text}"

        run = p.add_run(text)
        run.font.size = Pt(11)

        # Style based on level
        if entry.level == 1:
            run.font.bold = True

        body.insert(insert_position + inserted, p._element)
        inserted += 1

    # Add spacing
    space = doc.add_paragraph()
    body.insert(insert_position + inserted, space._element)
    inserted += 1

    return inserted


def insert_inferred_tof(
    doc: Document,
    entries: List[CaptionEntry],
    insert_position: int,
) -> int:
    """
    Insert an inferred Table of Figures as actual text.

    Returns number of paragraphs inserted.
    """
    if not entries:
        return 0

    body = doc.element.body
    inserted = 0

    # Add heading
    heading = doc.add_paragraph()
    heading.style = "Heading 1"
    heading.add_run("Figures")
    body.insert(insert_position, heading._element)
    inserted += 1

    # Add entries
    for entry in entries:
        p = doc.add_paragraph()

        # Format: "Figure X-Y: Caption text"
        if entry.caption:
            text = f"{entry.label}: {entry.caption}"
        else:
            text = entry.label

        run = p.add_run(text)
        run.font.size = Pt(11)

        # Apply TOF style if available
        try:
            p.style = "table of figures"
        except:
            pass

        body.insert(insert_position + inserted, p._element)
        inserted += 1

    # Add spacing
    space = doc.add_paragraph()
    body.insert(insert_position + inserted, space._element)
    inserted += 1

    return inserted


def insert_inferred_tot(
    doc: Document,
    entries: List[CaptionEntry],
    insert_position: int,
) -> int:
    """
    Insert an inferred Table of Tables as actual text.

    Returns number of paragraphs inserted.
    """
    if not entries:
        return 0

    body = doc.element.body
    inserted = 0

    # Add heading
    heading = doc.add_paragraph()
    heading.style = "Heading 1"
    heading.add_run("Tables")
    body.insert(insert_position, heading._element)
    inserted += 1

    # Add entries
    for entry in entries:
        p = doc.add_paragraph()

        # Format: "Table X-Y: Caption text"
        if entry.caption:
            text = f"{entry.label}: {entry.caption}"
        else:
            text = entry.label

        run = p.add_run(text)
        run.font.size = Pt(11)

        # Apply TOF style if available (same style for tables)
        try:
            p.style = "table of figures"
        except:
            pass

        body.insert(insert_position + inserted, p._element)
        inserted += 1

    # Add spacing
    space = doc.add_paragraph()
    body.insert(insert_position + inserted, space._element)
    inserted += 1

    return inserted


def create_toc_field(paragraph) -> None:
    """
    Insert a Table of Contents field into a paragraph.

    The TOC will need to be updated when opened in Word (Ctrl+A, F9).
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    # Add placeholder text
    placeholder_run = paragraph.add_run("(Update this field to generate Table of Contents: Select All → F9)")
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.font.italic = True

    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)


def create_tof_field(paragraph) -> None:
    """
    Insert a Table of Figures field into a paragraph.
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\h \\z \\c "Figure" '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    placeholder_run = paragraph.add_run("(Update field to generate Table of Figures)")
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.font.italic = True

    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)


def create_tot_field(paragraph) -> None:
    """
    Insert a Table of Tables field into a paragraph.
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\h \\z \\c "Table" '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    placeholder_run = paragraph.add_run("(Update field to generate Table of Tables)")
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.font.italic = True

    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)


def add_front_matter(
    doc: Document,
    analysis: DocumentAnalysis,
    config: RestructureConfig,
    insert_position: int = 0,
    inferred: Optional[InferredStructure] = None,
) -> int:
    """
    Add front matter (TOC, TOF, TOT) to document.

    If config.use_inferred_content is True and inferred is provided,
    inserts actual text content. Otherwise uses Word fields.

    Returns the number of paragraphs inserted.
    """
    inserted = 0
    use_inferred = config.use_inferred_content and inferred is not None

    if use_inferred:
        # Insert inferred content as actual text (in reverse order: TOT, TOF, TOC)
        if config.add_tot and inferred.table_entries:
            count = insert_inferred_tot(doc, inferred.table_entries, insert_position)
            inserted += count

        if config.add_tof and inferred.figure_entries:
            count = insert_inferred_tof(doc, inferred.figure_entries, insert_position)
            inserted += count

        if config.add_toc and inferred.toc_entries:
            count = insert_inferred_toc(doc, inferred.toc_entries, insert_position)
            inserted += count

    else:
        # Use Word fields (require manual update)
        elements_to_add = []

        if config.add_tot and analysis.table_count > 0:
            elements_to_add.append(("Tables", "tot"))

        if config.add_tof and analysis.figure_count > 0:
            elements_to_add.append(("Figures", "tof"))

        if config.add_toc:
            elements_to_add.append(("Contents", "toc"))

        if elements_to_add:
            for title, field_type in reversed(elements_to_add):
                heading_para = doc.add_paragraph()
                heading_para.style = "Heading 1"
                heading_para.add_run(title)

                body = doc.element.body
                body.insert(insert_position, heading_para._element)
                inserted += 1

                field_para = doc.add_paragraph()
                if field_type == "toc":
                    create_toc_field(field_para)
                elif field_type == "tof":
                    create_tof_field(field_para)
                elif field_type == "tot":
                    create_tot_field(field_para)

                body.insert(insert_position + 1, field_para._element)
                inserted += 1

                space_para = doc.add_paragraph()
                body.insert(insert_position + 2, space_para._element)
                inserted += 1

    return inserted


def add_missing_sections(
    doc: Document,
    analysis: DocumentAnalysis,
) -> List[str]:
    """
    Add placeholder sections for missing required elements.

    Returns list of sections added.
    """
    added = []

    # Add Abstract if missing
    if not analysis.abstract:
        # Find first heading and insert before it
        for i, p in enumerate(doc.paragraphs):
            if p.style and "heading" in p.style.name.lower():
                abstract_heading = doc.add_paragraph()
                abstract_heading.style = "Heading 1"
                abstract_heading.add_run("Abstract")

                abstract_content = doc.add_paragraph()
                run = abstract_content.add_run("[ABSTRACT CONTENT NEEDED]")
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

                # Move to position
                body = doc.element.body
                body.insert(i, abstract_heading._element)
                body.insert(i + 1, abstract_content._element)

                added.append("Abstract")
                break

    # Add References if missing (at end)
    if not analysis.has_references:
        ref_heading = doc.add_paragraph()
        ref_heading.style = "Heading 1"
        ref_heading.add_run("References")

        ref_content = doc.add_paragraph()
        run = ref_content.add_run("[REFERENCES NEEDED - Use APA format]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.bold = True

        added.append("References")

    return added


def add_field_update_instructions(doc: Document) -> None:
    """
    Add instructions at the top of the document for updating fields.
    """
    instruction_para = doc.add_paragraph()
    run = instruction_para.add_run(
        "⚠️ IMPORTANT: To update Table of Contents, Figures, and Tables: "
        "Select All (Ctrl+A) → Update Fields (F9) → Save"
    )
    run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
    run.font.bold = True
    run.font.size = Pt(11)

    # Move to very beginning
    body = doc.element.body
    body.insert(0, instruction_para._element)


def restructure_document(
    input_path: str,
    output_path: str,
    config: Optional[RestructureConfig] = None,
) -> Dict:
    """
    Restructure a document to follow DTC template.

    Args:
        input_path: Path to input DOCX
        output_path: Path for output DOCX
        config: Restructuring configuration

    Returns:
        Dictionary with restructuring statistics
    """
    config = config or RestructureConfig()
    doc = Document(input_path)

    # Analyze current structure
    analysis = analyze_document(doc)

    # Infer document structure for TOC/TOF/TOT content
    inferred = infer_document_structure(doc)

    stats = {
        "title": analysis.title,
        "original_figures": analysis.figure_count,
        "original_tables": analysis.table_count,
        "missing_elements": analysis.missing_elements.copy(),
        "elements_added": [],
        "requires_field_update": False,
        "inferred_toc_entries": len(inferred.toc_entries),
        "inferred_figure_entries": len(inferred.figure_entries),
        "inferred_table_entries": len(inferred.table_entries),
    }

    # Find where to insert front matter (after title, before content)
    insert_pos = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Title":
            insert_pos = i + 1
            break
        if p.text.strip():
            insert_pos = i
            break

    # Track what we're adding
    if not analysis.has_toc and config.add_toc:
        if config.use_inferred_content and inferred.toc_entries:
            stats["elements_added"].append(f"Table of Contents ({len(inferred.toc_entries)} entries)")
        else:
            stats["elements_added"].append("Table of Contents (field)")
            stats["requires_field_update"] = True

    if not analysis.has_tof and config.add_tof:
        if inferred.figure_entries:
            if config.use_inferred_content:
                stats["elements_added"].append(f"Table of Figures ({len(inferred.figure_entries)} entries)")
            else:
                stats["elements_added"].append("Table of Figures (field)")
                stats["requires_field_update"] = True
        elif analysis.figure_count > 0:
            stats["elements_added"].append("Table of Figures (field - no captions detected)")
            stats["requires_field_update"] = True

    if not analysis.has_tot and config.add_tot:
        if inferred.table_entries:
            if config.use_inferred_content:
                stats["elements_added"].append(f"Table of Tables ({len(inferred.table_entries)} entries)")
            else:
                stats["elements_added"].append("Table of Tables (field)")
                stats["requires_field_update"] = True
        elif analysis.table_count > 0:
            stats["elements_added"].append("Table of Tables (field - no captions detected)")
            stats["requires_field_update"] = True

    # Add front matter with inferred content
    add_front_matter(doc, analysis, config, insert_pos, inferred)

    # Add missing sections if configured
    if config.include_missing_placeholders:
        sections_added = add_missing_sections(doc, analysis)
        stats["elements_added"].extend(sections_added)

    # Add field update instructions only if needed
    if stats["requires_field_update"] and not config.use_inferred_content:
        add_field_update_instructions(doc)

    # Save
    doc.save(output_path)

    return stats


def check_template_compliance(doc_path: str) -> Dict:
    """
    Check a document's compliance with DTC template requirements.

    Returns a compliance report.
    """
    doc = Document(doc_path)
    analysis = analyze_document(doc)

    compliance = {
        "compliant": len(analysis.missing_elements) == 0,
        "title": analysis.title,
        "has_toc": analysis.has_toc,
        "has_tof": analysis.has_tof or analysis.figure_count == 0,
        "has_tot": analysis.has_tot or analysis.table_count == 0,
        "has_abstract": analysis.abstract is not None,
        "has_references": analysis.has_references,
        "figure_count": analysis.figure_count,
        "table_count": analysis.table_count,
        "section_count": len(analysis.sections),
        "missing_elements": analysis.missing_elements,
        "recommendations": [],
    }

    # Generate recommendations
    if not analysis.has_toc:
        compliance["recommendations"].append(
            "Add Table of Contents after title page"
        )
    if analysis.figure_count > 0 and not analysis.has_tof:
        compliance["recommendations"].append(
            f"Add Table of Figures ({analysis.figure_count} figures detected)"
        )
    if analysis.table_count > 0 and not analysis.has_tot:
        compliance["recommendations"].append(
            f"Add Table of Tables ({analysis.table_count} tables detected)"
        )
    if not analysis.abstract:
        compliance["recommendations"].append(
            "Add Abstract or Executive Summary section"
        )
    if not analysis.has_references:
        compliance["recommendations"].append(
            "Add References section with APA-formatted citations"
        )

    return compliance
