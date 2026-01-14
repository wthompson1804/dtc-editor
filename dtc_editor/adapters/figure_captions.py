"""
Figure Caption Generator for DTC Editorial Engine.

Detects figures without captions and either:
1. Infers caption from surrounding context using LLM
2. Adds placeholder "SOURCE AND/OR DESCRIPTION NEEDED" in red font
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


@dataclass
class FigureInfo:
    """Information about a figure in the document."""
    para_index: int
    section_number: str  # e.g., "3" for section 3
    figure_number: int   # sequential within section
    context_before: str
    context_after: str
    has_caption: bool = False
    inferred_caption: Optional[str] = None
    needs_placeholder: bool = False


@dataclass
class CaptionConfig:
    """Configuration for caption generation."""
    use_llm: bool = True
    api_key: Optional[str] = None
    model: str = "claude-sonnet-4-20250514"
    placeholder_text: str = "SOURCE AND/OR DESCRIPTION NEEDED"
    placeholder_color: Tuple[int, int, int] = (255, 0, 0)  # Red


def detect_figures(doc: Document) -> List[FigureInfo]:
    """
    Detect all figures (drawings) in the document.

    Returns list of FigureInfo with context for caption inference.
    """
    figures = []
    current_section = "1"
    section_fig_count = {}

    for i, para in enumerate(doc.paragraphs):
        # Track section numbers from headings
        if para.style and para.style.name.lower().startswith("heading"):
            # Try to extract section number from heading
            match = re.match(r'^(\d+(?:\.\d+)*)', para.text.strip())
            if match:
                current_section = match.group(1).split('.')[0]
            else:
                # Increment section on major headings without numbers
                if "heading 1" in para.style.name.lower():
                    try:
                        current_section = str(int(current_section) + 1)
                    except ValueError:
                        pass

        # Check for drawings/figures
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            # Initialize section counter
            if current_section not in section_fig_count:
                section_fig_count[current_section] = 0
            section_fig_count[current_section] += 1

            # Get context
            context_before = ""
            context_after = ""

            # Look back for context (up to 3 paragraphs)
            for j in range(max(0, i - 3), i):
                txt = doc.paragraphs[j].text.strip()
                if txt:
                    context_before = txt
                    break

            # Look ahead for context (up to 3 paragraphs)
            for j in range(i + 1, min(len(doc.paragraphs), i + 4)):
                txt = doc.paragraphs[j].text.strip()
                if txt:
                    context_after = txt
                    break

            # Check if next paragraph looks like a caption
            has_caption = False
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip().lower()
                if next_text.startswith("figure ") and re.match(r'^figure\s+\d', next_text):
                    has_caption = True

            figures.append(FigureInfo(
                para_index=i,
                section_number=current_section,
                figure_number=section_fig_count[current_section],
                context_before=context_before,
                context_after=context_after,
                has_caption=has_caption,
            ))

    return figures


def infer_caption_from_context(
    figure: FigureInfo,
    api_key: Optional[str] = None,
    model: str = "claude-sonnet-4-20250514",
) -> Optional[str]:
    """
    Use LLM to infer a figure caption from surrounding context.

    Returns inferred caption or None if can't be determined.
    """
    if not api_key:
        return None

    try:
        from anthropic import Anthropic
    except ImportError:
        return None

    client = Anthropic(api_key=api_key)

    prompt = f"""Based on the surrounding text context, generate a brief, descriptive caption for a figure in a technical document.

Text BEFORE the figure:
"{figure.context_before}"

Text AFTER the figure:
"{figure.context_after}"

Guidelines:
- Caption should be 5-15 words describing what the figure shows
- Use sentence case (only capitalize first word and proper nouns)
- Do not include "Figure X-Y:" prefix - just the caption text
- Focus on WHAT is shown, not interpretation
- If the context clearly describes the figure (e.g., "shown in the following diagram"), use that description
- If you cannot determine what the figure shows with reasonable confidence, respond with exactly: CANNOT_INFER

Respond with ONLY the caption text or CANNOT_INFER."""

    try:
        response = client.messages.create(
            model=model,
            max_tokens=100,
            temperature=0.3,
            messages=[{"role": "user", "content": prompt}]
        )

        caption = response.content[0].text.strip()

        if caption == "CANNOT_INFER" or len(caption) < 5:
            return None

        # Clean up caption
        caption = caption.rstrip('.')
        # Remove any accidentally included figure prefix
        caption = re.sub(r'^figure\s+\d+[-.]?\d*:?\s*', '', caption, flags=re.IGNORECASE)

        return caption

    except Exception as e:
        print(f"Warning: LLM caption inference failed: {e}")
        return None


def generate_figure_label(section: str, fig_num: int) -> str:
    """Generate DTC-style figure label: Figure X-Y"""
    return f"Figure {section}-{fig_num}"


def add_caption_to_document(
    doc: Document,
    figure: FigureInfo,
    caption_text: str,
    is_placeholder: bool = False,
    placeholder_color: Tuple[int, int, int] = (255, 0, 0),
) -> None:
    """
    Add a caption paragraph after the figure.

    For placeholders, uses red font color.
    """
    # Generate the full caption
    label = generate_figure_label(figure.section_number, figure.figure_number)

    if is_placeholder:
        full_caption = f"{label}: {caption_text}"
    else:
        # Ensure caption ends with period
        if not caption_text.endswith('.'):
            caption_text += '.'
        full_caption = f"{label}: {caption_text}"

    # Insert new paragraph after the figure
    # We need to insert after figure.para_index
    fig_para = doc.paragraphs[figure.para_index]

    # Create new paragraph element
    new_p = doc.add_paragraph()

    # Move it to the correct position (after the figure)
    fig_para._element.addnext(new_p._element)

    # Style the caption
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the text with appropriate formatting
    run = new_p.add_run(full_caption)
    run.font.size = Pt(10)
    run.font.italic = True

    if is_placeholder:
        run.font.color.rgb = RGBColor(*placeholder_color)
        run.font.bold = True


def process_figure_captions(
    docx_path: str,
    output_path: str,
    config: CaptionConfig,
) -> dict:
    """
    Process a document to add missing figure captions.

    Returns statistics about the processing.
    """
    doc = Document(docx_path)
    figures = detect_figures(doc)

    stats = {
        "total_figures": len(figures),
        "already_captioned": 0,
        "captions_inferred": 0,
        "placeholders_added": 0,
        "figures_processed": [],
    }

    # Process figures in reverse order to avoid index shifting
    for figure in reversed(figures):
        if figure.has_caption:
            stats["already_captioned"] += 1
            continue

        # Try to infer caption using LLM
        inferred = None
        if config.use_llm and config.api_key:
            inferred = infer_caption_from_context(
                figure,
                api_key=config.api_key,
                model=config.model,
            )

        if inferred:
            add_caption_to_document(
                doc, figure, inferred,
                is_placeholder=False,
            )
            stats["captions_inferred"] += 1
            stats["figures_processed"].append({
                "index": figure.para_index,
                "label": generate_figure_label(figure.section_number, figure.figure_number),
                "caption": inferred,
                "type": "inferred",
            })
        else:
            add_caption_to_document(
                doc, figure, config.placeholder_text,
                is_placeholder=True,
                placeholder_color=config.placeholder_color,
            )
            stats["placeholders_added"] += 1
            stats["figures_processed"].append({
                "index": figure.para_index,
                "label": generate_figure_label(figure.section_number, figure.figure_number),
                "caption": config.placeholder_text,
                "type": "placeholder",
            })

    doc.save(output_path)
    return stats


def add_figure_captions_to_ir(
    docx_path: str,
    output_path: str,
    use_llm: bool = True,
    api_key: Optional[str] = None,
) -> dict:
    """
    Convenience function to add figure captions to a document.

    Args:
        docx_path: Path to input DOCX
        output_path: Path to save output DOCX
        use_llm: Whether to use LLM for caption inference
        api_key: Anthropic API key (required if use_llm=True)

    Returns:
        Statistics dictionary
    """
    config = CaptionConfig(
        use_llm=use_llm,
        api_key=api_key,
    )

    return process_figure_captions(docx_path, output_path, config)
