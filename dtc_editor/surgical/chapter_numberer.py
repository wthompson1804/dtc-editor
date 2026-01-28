"""
Chapter/Section Numberer for Surgical Pipeline

Handles:
1. Detection of ALL heading levels (Heading 1, 2, 3, etc.)
2. Stripping ALL existing numbers from headings
3. Applying consistent hierarchical numbering (1, 1.1, 1.1.1, 2, 2.1, etc.)
4. Preserving special sections without numbers (Abstract, References, etc.)

Reference: dtc_editor/rules/surgical_rules_manifest.yml (structure.chapters.numbered)
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Set, Optional, Dict, Tuple
from pathlib import Path
import re
import logging

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class HeadingInfo:
    """Information about a heading at any level."""
    para_index: int
    level: int  # 1, 2, 3, etc.
    original_text: str
    existing_number: Optional[str]  # "1", "1.2", "1.2.3" etc. if numbered
    title_text: str  # Text after number (or full text if unnumbered)
    is_special: bool  # True for Abstract, References, etc.
    assigned_number: Optional[str]  # New hierarchical number to assign ("1", "1.1", etc.)


@dataclass
class ChapterNumbererConfig:
    """Configuration for chapter/section numbering."""
    # Special sections that don't get numbers (checked against title text)
    unnumbered_sections: Set[str] = field(default_factory=lambda: {
        # Front matter / Title page content
        "abstract",
        "executive summary",
        "white paper",
        "technical paper",
        "working paper",
        "position paper",
        # Back matter
        "references",
        "authors",
        "authors & legal notice",
        "appendix",
        "annex",
        "acknowledgments",
        "acknowledgements",
        "glossary",
        "table of contents",
        "toc",
        "bibliography",
        "index",
        # Common document titles
        "consortium",
        "preface",
        "foreword",
        "figures",
        "tables",
    })

    # Maximum heading level to number (1=H1 only, 2=H1+H2, 3=H1+H2+H3, etc.)
    max_level: int = 3

    # Whether to add separator between number and title
    separator: str = " "  # Space between "1.2" and "Section Title"


@dataclass
class ChapterNumbererResult:
    """Result of chapter/section numbering."""
    headings_found: int
    headings_numbered: int
    headings_renumbered: int
    special_sections: int
    changes: List[Dict]
    issues: List[str]
    # For backwards compatibility
    @property
    def chapters_found(self) -> int:
        return self.headings_found
    @property
    def chapters_numbered(self) -> int:
        return self.headings_numbered
    @property
    def chapters_renumbered(self) -> int:
        return self.headings_renumbered


# =============================================================================
# Main Processor Class
# =============================================================================

class ChapterNumberer:
    """
    Adds/corrects hierarchical section numbers in a DOCX document.

    Handles ALL heading levels with proper hierarchical numbering:
    - Heading 1 → 1, 2, 3, ...
    - Heading 2 → 1.1, 1.2, 2.1, ...
    - Heading 3 → 1.1.1, 1.1.2, 2.1.1, ...

    Workflow:
    1. Scan document for ALL heading styles
    2. Strip ALL existing numbers from headings
    3. Identify special sections (Abstract, References, etc.)
    4. Assign hierarchical numbers to content sections
    5. Rewrite all headings with clean formatting
    """

    def __init__(self, config: ChapterNumbererConfig = None):
        self.config = config or ChapterNumbererConfig()
        self.headings: List[HeadingInfo] = []
        self.doc: Optional[Document] = None

    def process(self, doc: Document) -> ChapterNumbererResult:
        """Process all headings in the document."""
        self.doc = doc
        self.headings = []

        changes = []
        issues = []

        # Step 1: Scan ALL headings
        self._scan_all_headings()
        logger.info(f"Found {len(self.headings)} headings")

        # Step 2: Assign hierarchical numbers
        self._assign_hierarchical_numbers()

        # Step 3: Apply changes - ALWAYS rewrite to ensure clean formatting
        headings_numbered = 0
        headings_renumbered = 0
        special_sections = 0

        for heading in self.headings:
            if heading.is_special:
                special_sections += 1
                # Still rewrite special sections to strip any stray numbers
                self._rewrite_heading_without_number(heading)
                continue

            if heading.assigned_number is None:
                continue

            # ALWAYS rewrite the heading
            self._rewrite_heading(heading)

            if heading.existing_number is None:
                headings_numbered += 1
                changes.append({
                    "type": "heading_numbered",
                    "para_index": heading.para_index,
                    "level": heading.level,
                    "title": heading.title_text,
                    "new_number": heading.assigned_number,
                })
            else:
                headings_renumbered += 1
                changes.append({
                    "type": "heading_renumbered",
                    "para_index": heading.para_index,
                    "level": heading.level,
                    "title": heading.title_text,
                    "old_number": heading.existing_number,
                    "new_number": heading.assigned_number,
                })

        return ChapterNumbererResult(
            headings_found=len(self.headings),
            headings_numbered=headings_numbered,
            headings_renumbered=headings_renumbered,
            special_sections=special_sections,
            changes=changes,
            issues=issues,
        )

    # =========================================================================
    # Scanning Methods
    # =========================================================================

    def _get_heading_level(self, style_name: str) -> Optional[int]:
        """Extract heading level from style name."""
        style_lower = style_name.lower()

        # Match "heading 1", "heading 2", etc.
        match = re.search(r'heading\s*(\d+)', style_lower)
        if match:
            level = int(match.group(1))
            if 1 <= level <= self.config.max_level:
                return level
        return None

    def _strip_leading_numbers(self, text: str) -> Tuple[Optional[str], str]:
        """
        Strip ALL leading numbers from text.

        Returns:
            (existing_number, clean_title)

        Handles patterns like:
        - "1 Title" → ("1", "Title")
        - "1.2 Title" → ("1.2", "Title")
        - "1.2.3 Title" → ("1.2.3", "Title")
        - "4    4 Title" (corrupted) → ("4", "Title")
        - "1. Title" → ("1.", "Title")
        """
        existing_number = None
        title_text = text.strip()

        # Loop to strip ALL leading number patterns
        while True:
            # Match: digit(s), optional decimal parts, optional trailing dot, optional whitespace
            match = re.match(r'^(\d+(?:\.\d+)*\.?)\s*(.*)$', title_text)
            if match and match.group(2).strip():
                # Found a leading number - save first one found
                if existing_number is None:
                    existing_number = match.group(1)
                title_text = match.group(2).strip()
            else:
                break

        # Safety check: if we stripped everything, restore original
        if not title_text.strip():
            return None, text.strip()

        return existing_number, title_text

    def _scan_all_headings(self) -> None:
        """Scan document for all heading styles up to max_level."""
        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else ""
            level = self._get_heading_level(style_name)

            if level is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            # Strip existing numbers
            existing_number, title_text = self._strip_leading_numbers(text)

            # Check if this is a special section
            title_lower = title_text.lower().strip()
            is_special = any(
                special in title_lower or title_lower in special
                for special in self.config.unnumbered_sections
            )

            self.headings.append(HeadingInfo(
                para_index=i,
                level=level,
                original_text=text,
                existing_number=existing_number,
                title_text=title_text,
                is_special=is_special,
                assigned_number=None,
            ))

    def _assign_hierarchical_numbers(self) -> None:
        """Assign hierarchical numbers to non-special headings."""
        # Counters for each level
        counters = [0] * (self.config.max_level + 1)  # Index 0 unused, 1-based

        for heading in self.headings:
            if heading.is_special:
                heading.assigned_number = None
                continue

            level = heading.level

            # Increment counter at this level
            counters[level] += 1

            # Reset all counters below this level
            for l in range(level + 1, len(counters)):
                counters[l] = 0

            # Build hierarchical number string
            number_parts = [str(counters[l]) for l in range(1, level + 1)]
            heading.assigned_number = ".".join(number_parts)

    # =========================================================================
    # Modification Methods
    # =========================================================================

    def _rewrite_heading(self, heading: HeadingInfo) -> bool:
        """Rewrite heading with proper number and formatting."""
        if heading.assigned_number is None:
            return False

        para = self.doc.paragraphs[heading.para_index]

        # Build new text: "1.2 Section Title"
        clean_title = heading.title_text.strip()
        new_text = f"{heading.assigned_number}{self.config.separator}{clean_title}"

        self._update_paragraph_text(para, new_text)
        logger.info(f"Rewrote H{heading.level}: '{heading.original_text}' → '{new_text}'")
        return True

    def _rewrite_heading_without_number(self, heading: HeadingInfo) -> bool:
        """Rewrite special heading to strip any existing number."""
        para = self.doc.paragraphs[heading.para_index]

        # Just use the clean title text (numbers already stripped)
        clean_title = heading.title_text.strip()

        # Only rewrite if there was an existing number to strip
        if heading.existing_number:
            self._update_paragraph_text(para, clean_title)
            logger.info(f"Stripped number from special heading: '{heading.original_text}' → '{clean_title}'")
        return True

    def _update_paragraph_text(self, para, new_text: str) -> None:
        """Update paragraph text while preserving formatting."""
        # Get formatting from first run if available
        first_run = para.runs[0] if para.runs else None
        font_name = None
        font_size = None
        font_bold = None

        if first_run:
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold

        # Clear and rewrite
        para.clear()
        run = para.add_run(new_text)

        # Restore formatting
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold


# =============================================================================
# Convenience Functions
# =============================================================================

def number_chapters(
    docx_path: str,
    output_path: str,
    config: Optional[ChapterNumbererConfig] = None,
) -> ChapterNumbererResult:
    """Add/correct section numbers in a DOCX file."""
    if config is None:
        config = ChapterNumbererConfig()

    doc = Document(docx_path)
    processor = ChapterNumberer(config)
    result = processor.process(doc)

    doc.save(output_path)
    return result


def analyze_chapters(docx_path: str) -> List[Dict]:
    """Analyze heading structure without making changes."""
    doc = Document(docx_path)
    config = ChapterNumbererConfig()
    processor = ChapterNumberer(config)

    processor.doc = doc
    processor._scan_all_headings()
    processor._assign_hierarchical_numbers()

    return [
        {
            "para_index": h.para_index,
            "level": h.level,
            "original": h.original_text,
            "existing_number": h.existing_number,
            "title": h.title_text,
            "is_special": h.is_special,
            "assigned_number": h.assigned_number,
        }
        for h in processor.headings
    ]
