"""
Acronym Expander for Surgical Pipeline

Handles:
1. Detection of unexpanded acronyms on first use
2. Auto-expansion using known acronym database
3. LLM lookup for unknown acronyms (optional)
4. Organization acronyms that don't need expansion

Reference: dtc_editor/rules/surgical_rules_manifest.yml (acronyms section)
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Set, Dict, Optional, Tuple
from pathlib import Path
import re
import logging

from docx import Document
from docx.shared import Pt

# Import known acronyms from holistic module
from dtc_editor.holistic.acronyms import (
    EXPANDABLE_ACRONYMS,
    ORGANIZATION_ACRONYMS,
    get_expansion,
    format_first_use,
    is_organization_acronym,
)

logger = logging.getLogger(__name__)


# =============================================================================
# Additional Domain-Specific Acronyms (DTC/Energy)
# =============================================================================

# Extend with DTC/energy domain acronyms
DTC_ACRONYMS: Dict[str, str] = {
    "VPP": "Virtual Power Plant",
    "DER": "Distributed Energy Resource",
    "DERMS": "Distributed Energy Resource Management System",
    "EMS": "Energy Management System",
    "SCADA": "Supervisory Control and Data Acquisition",
    "BMS": "Building Management System",
    "BEMS": "Building Energy Management System",
    "PV": "Photovoltaic",
    "EV": "Electric Vehicle",
    "V2G": "Vehicle-to-Grid",
    "V2X": "Vehicle-to-Everything",
    "ESS": "Energy Storage System",
    "BESS": "Battery Energy Storage System",
    "DR": "Demand Response",
    "DSM": "Demand Side Management",
    "AMI": "Advanced Metering Infrastructure",
    "ADR": "Automated Demand Response",
    "TSO": "Transmission System Operator",
    "DSO": "Distribution System Operator",
    "ISO": "Independent System Operator",  # Note: Also org acronym in different context
    "RTO": "Regional Transmission Organization",
    "FERC": "Federal Energy Regulatory Commission",
    "NERC": "North American Electric Reliability Corporation",
    "CIM": "Common Information Model",
    "IEC 61850": "IEC 61850",  # Standard reference, keep as-is
    "OpenADR": "Open Automated Demand Response",
    "OPC UA": "Open Platform Communications Unified Architecture",
    # Digital Twin specific
    "DTaaS": "Digital Twin as a Service",
    "CPS": "Cyber-Physical System",
    "M2M": "Machine-to-Machine",
    "CMMS": "Computerized Maintenance Management System",
    "APM": "Asset Performance Management",
    "BIM": "Building Information Modeling",
    "GIS": "Geographic Information System",
    # Telecom/Edge
    "UPF": "User Plane Function",
    "SMF": "Session Management Function",
    "AMF": "Access and Mobility Management Function",
    "NFVI": "Network Functions Virtualization Infrastructure",
    "MANO": "Management and Network Orchestration",
    "ONAP": "Open Network Automation Platform",
}

# Combine all known acronyms
ALL_ACRONYMS: Dict[str, str] = {**EXPANDABLE_ACRONYMS, **DTC_ACRONYMS}


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class AcronymOccurrence:
    """A single occurrence of an acronym in the document."""
    para_index: int
    span_start: int
    span_end: int
    acronym: str
    is_first_use: bool
    is_already_expanded: bool
    context: str  # Surrounding text for context


@dataclass
class AcronymExpanderConfig:
    """Configuration for acronym expansion."""
    # Known acronyms database
    known_acronyms: Dict[str, str] = field(default_factory=lambda: ALL_ACRONYMS.copy())
    organization_acronyms: Set[str] = field(default_factory=lambda: ORGANIZATION_ACRONYMS.copy())

    # False positives to ignore (template placeholders, common words, etc.)
    ignore_patterns: Set[str] = field(default_factory=lambda: {
        "DRAFT", "STATE", "YYYY", "MM", "DD", "TBD", "TODO", "FIXME",
        "NOTE", "NB", "IE", "EG", "PS", "AM", "PM", "USA", "UK", "EU",
        "BC", "AD", "CEO", "CFO", "CTO", "COO", "VP", "HR", "PR",
        "OK", "ID", "VS", "NO", "YES", "NA", "TBA",
    })

    # Behavior - LLM lookup is now ENABLED by default
    expand_unknown: bool = True  # Expand unknown acronyms using LLM
    use_llm_lookup: bool = True  # Use LLM for unknown acronyms
    confidence_threshold: float = 0.7  # Lower threshold - expand even with uncertainty
    max_llm_lookups: int = 50  # Increased limit for LLM calls per document

    # API key for LLM lookups (optional - can be set via env var)
    anthropic_api_key: Optional[str] = None

    # Format
    expansion_format: str = "{expansion} ({acronym})"

    # Paragraphs to skip (TOC, captions, etc.)
    skip_styles: Set[str] = field(default_factory=lambda: {
        "toc", "table of contents", "caption", "heading"
    })


@dataclass
class AcronymExpanderResult:
    """Result of acronym expansion."""
    acronyms_found: int
    expansions_made: int
    unknown_acronyms: List[str]
    already_expanded: int
    organization_acronyms_skipped: int
    changes: List[Dict]
    issues: List[str]
    # NEW: Track items that need human review
    requires_review: List[Dict] = field(default_factory=list)  # LLM-expanded items needing review


# =============================================================================
# Main Processor Class
# =============================================================================

class AcronymExpander:
    """
    Expands acronyms on first use in a DOCX document.

    Workflow:
    1. Scan document for all acronyms (ALL_CAPS patterns)
    2. Track which have been expanded (already have "Expansion (ACR)" format)
    3. Find first occurrence of each unexpanded acronym
    4. Expand with "Full Name (ACRONYM)" format
    5. For unknown acronyms, use LLM lookup and flag for review
    """

    def __init__(self, config: AcronymExpanderConfig, llm_client=None):
        self.config = config
        self.occurrences: List[AcronymOccurrence] = []
        self.defined: Set[str] = set()  # Acronyms already defined in document
        self.doc: Optional[Document] = None
        self._anthropic_client = None

        # Set up LLM client if API key is available
        if llm_client:
            self.llm_client = llm_client
        elif config.use_llm_lookup:
            self.llm_client = self._create_llm_client()
        else:
            self.llm_client = None

    def _create_llm_client(self):
        """Create an Anthropic client for LLM lookups."""
        import os
        api_key = self.config.anthropic_api_key or os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            logger.warning("No Anthropic API key found. LLM acronym lookup disabled.")
            return None

        try:
            import anthropic
            self._anthropic_client = anthropic.Anthropic(api_key=api_key)
            return self._anthropic_client
        except ImportError:
            logger.warning("anthropic library not installed. LLM acronym lookup disabled.")
            return None
        except Exception as e:
            logger.warning(f"Failed to create Anthropic client: {e}")
            return None

    def process(self, doc: Document) -> AcronymExpanderResult:
        """
        Process all acronyms in the document.

        Args:
            doc: python-docx Document object

        Returns:
            AcronymExpanderResult with statistics and changes
        """
        self.doc = doc
        self.occurrences = []
        self.defined = set()

        changes = []
        issues = []
        unknown_acronyms = []
        requires_review = []

        # Step 1: Scan for existing expansions (already defined)
        self._scan_existing_expansions()
        logger.info(f"Found {len(self.defined)} already-expanded acronyms")

        # Step 2: Scan for all acronym occurrences
        self._scan_acronyms()
        logger.info(f"Found {len(self.occurrences)} acronym occurrences")

        # Step 3: Identify first uses that need expansion
        first_uses = self._identify_first_uses()
        logger.info(f"Found {len(first_uses)} acronyms needing expansion")

        # Step 4: Expand each first use
        expansions_made = 0
        org_skipped = 0
        llm_lookups = 0

        # Process in reverse order to avoid index shifting
        for occurrence in reversed(first_uses):
            acronym = occurrence.acronym

            # Skip organization acronyms
            if acronym in self.config.organization_acronyms:
                org_skipped += 1
                continue

            # Get expansion from known database
            expansion = self.config.known_acronyms.get(acronym)
            is_llm_expansion = False

            # If not in known database, try LLM lookup
            if not expansion and self.config.use_llm_lookup and self.llm_client:
                if llm_lookups < self.config.max_llm_lookups:
                    llm_result = self._lookup_acronym_llm(acronym, occurrence.context)
                    llm_lookups += 1

                    if llm_result:
                        expansion = llm_result.get("expansion")
                        is_llm_expansion = True
                        # Add to known acronyms for this session
                        if expansion:
                            self.config.known_acronyms[acronym] = expansion

            if expansion:
                # Do the expansion
                success = self._expand_acronym(occurrence, expansion)
                if success:
                    expansions_made += 1

                    change_record = {
                        "type": "acronym_expanded_llm" if is_llm_expansion else "acronym_expanded",
                        "acronym": acronym,
                        "expansion": expansion,
                        "para_index": occurrence.para_index,
                        "requires_review": is_llm_expansion,
                    }
                    changes.append(change_record)

                    # LLM expansions always need human review
                    if is_llm_expansion:
                        requires_review.append({
                            "type": "llm_acronym_expansion",
                            "acronym": acronym,
                            "expansion": expansion,
                            "context": occurrence.context,
                            "para_index": occurrence.para_index,
                            "reason": "LLM-determined expansion - verify correctness",
                        })
            else:
                # Still unknown after LLM attempt
                if acronym not in unknown_acronyms:
                    unknown_acronyms.append(acronym)
                    issues.append(f"Unknown acronym: {acronym} (para {occurrence.para_index})")

        return AcronymExpanderResult(
            acronyms_found=len(self.occurrences),
            expansions_made=expansions_made,
            unknown_acronyms=unknown_acronyms,
            already_expanded=len(self.defined),
            organization_acronyms_skipped=org_skipped,
            changes=changes,
            issues=issues,
            requires_review=requires_review,
        )

    # =========================================================================
    # Scanning Methods
    # =========================================================================

    def _scan_existing_expansions(self) -> None:
        """Scan document for already-expanded acronyms."""
        full_text = "\n".join(p.text for p in self.doc.paragraphs)

        for acronym, expansion in self.config.known_acronyms.items():
            # Pattern: "Full Expansion (ACRONYM)"
            pattern = re.escape(expansion) + r'\s*\(' + re.escape(acronym) + r'\)'
            if re.search(pattern, full_text, re.IGNORECASE):
                self.defined.add(acronym)

    def _scan_acronyms(self) -> None:
        """Scan document for all acronym occurrences."""
        # Pattern for acronyms: 2+ uppercase letters, possibly with numbers
        # Examples: IoT, API, 5G, IEC, V2G, CI/CD
        acronym_pattern = re.compile(
            r'\b([A-Z][A-Za-z0-9/]*[A-Z][A-Za-z0-9]*|[A-Z]{2,})\b'
        )

        for i, para in enumerate(self.doc.paragraphs):
            # Skip certain paragraph styles
            style_name = para.style.name.lower() if para.style else ""
            if any(skip in style_name for skip in self.config.skip_styles):
                continue

            text = para.text

            for match in acronym_pattern.finditer(text):
                acronym = match.group(1)

                # Skip ignored patterns (false positives)
                if acronym in self.config.ignore_patterns:
                    continue

                # Skip if not a known acronym and not in our detection list
                if (acronym not in self.config.known_acronyms and
                    acronym not in self.config.organization_acronyms):
                    # Could be a real acronym we don't know about
                    # Only flag if it's all caps and 2-6 characters
                    if not (acronym.isupper() and 2 <= len(acronym) <= 6):
                        continue

                # Get context
                start = max(0, match.start() - 30)
                end = min(len(text), match.end() + 30)
                context = text[start:end]

                self.occurrences.append(AcronymOccurrence(
                    para_index=i,
                    span_start=match.start(),
                    span_end=match.end(),
                    acronym=acronym,
                    is_first_use=False,  # Will be determined later
                    is_already_expanded=self._check_if_expanded_here(text, match, acronym),
                    context=context,
                ))

    def _check_if_expanded_here(self, text: str, match: re.Match, acronym: str) -> bool:
        """Check if this specific occurrence is part of an expansion."""
        # Check if preceded by the expansion text
        expansion = self.config.known_acronyms.get(acronym, "")
        if expansion:
            # Look for "Expansion (" before the match
            before_start = max(0, match.start() - len(expansion) - 5)
            before_text = text[before_start:match.start()]
            if expansion.lower() in before_text.lower() and "(" in before_text:
                return True
        return False

    def _identify_first_uses(self) -> List[AcronymOccurrence]:
        """Identify first use of each unexpanded acronym."""
        seen: Set[str] = set(self.defined)  # Start with already-defined
        first_uses = []

        for occ in self.occurrences:
            if occ.is_already_expanded:
                seen.add(occ.acronym)
                continue

            if occ.acronym not in seen:
                occ.is_first_use = True
                first_uses.append(occ)
                seen.add(occ.acronym)

        return first_uses

    def _lookup_acronym_llm(self, acronym: str, context: str) -> Optional[Dict]:
        """
        Look up an unknown acronym using LLM.

        Returns:
            Dict with 'expansion' key if found, None otherwise.
            All LLM expansions are flagged for human review.
        """
        if not self.llm_client:
            return None

        prompt = f"""What does the acronym "{acronym}" stand for in this technical context?

Context: "{context}"

Instructions:
1. If you know the expansion with reasonable confidence, respond with ONLY the expansion (e.g., "Application Programming Interface").
2. If you're completely unsure, respond with "UNKNOWN".
3. Do not include the acronym in parentheses, just the expansion.
4. Prefer technical/industry-standard meanings over generic ones.
5. If there are multiple possible meanings, choose the most likely one for this technical context."""

        try:
            # Use the Anthropic client directly
            message = self.llm_client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=100,
                temperature=0.2,  # Low temperature for factual lookups
                messages=[{"role": "user", "content": prompt}],
            )

            # Extract response text
            expansion = ""
            for block in message.content:
                if hasattr(block, "text"):
                    expansion += block.text
            expansion = expansion.strip()

            # Validate response
            if expansion.upper() == "UNKNOWN" or len(expansion) < 3:
                return None

            # Basic sanity check - expansion should be longer than acronym
            if len(expansion) <= len(acronym):
                return None

            # Clean up any quotes or extra formatting
            expansion = expansion.strip('"\'')

            logger.info(f"LLM lookup: {acronym} → {expansion}")
            return {"expansion": expansion}

        except Exception as e:
            logger.warning(f"LLM lookup failed for {acronym}: {e}")
            return None

    # =========================================================================
    # Modification Methods
    # =========================================================================

    def _expand_acronym(self, occurrence: AcronymOccurrence, expansion: str) -> bool:
        """Expand an acronym at the given occurrence."""
        para = self.doc.paragraphs[occurrence.para_index]
        text = para.text

        # Build the expanded form
        # Per AP style, the expansion is lowercase (it's a generic term, not a proper noun)
        # The acronym remains uppercase in parentheses
        expansion_lower = expansion.lower()
        expanded = self.config.expansion_format.format(
            expansion=expansion_lower,
            acronym=occurrence.acronym
        )

        # Replace the acronym with expanded form
        new_text = (
            text[:occurrence.span_start] +
            expanded +
            text[occurrence.span_end:]
        )

        # Update paragraph (preserve formatting from first run)
        self._update_paragraph_text(para, new_text)

        logger.info(f"Expanded: {occurrence.acronym} → {expanded}")
        return True

    def _update_paragraph_text(self, para, new_text: str) -> None:
        """Update paragraph text while trying to preserve formatting."""
        # Get formatting from first run if available
        first_run = para.runs[0] if para.runs else None
        font_name = None
        font_size = None

        if first_run:
            font_name = first_run.font.name
            font_size = first_run.font.size

        # Clear and rewrite
        para.clear()
        run = para.add_run(new_text)

        # Restore formatting
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size


# =============================================================================
# Convenience Functions
# =============================================================================

def expand_acronyms(
    docx_path: str,
    output_path: str,
    config: Optional[AcronymExpanderConfig] = None,
) -> AcronymExpanderResult:
    """
    Expand acronyms on first use in a DOCX file.

    Args:
        docx_path: Path to input DOCX
        output_path: Path to save output DOCX
        config: Optional configuration

    Returns:
        AcronymExpanderResult with statistics
    """
    if config is None:
        config = AcronymExpanderConfig()

    doc = Document(docx_path)
    processor = AcronymExpander(config)
    result = processor.process(doc)

    doc.save(output_path)

    return result


def analyze_acronyms(docx_path: str) -> Dict:
    """
    Analyze acronym usage without making changes.

    Returns dict with acronym statistics.
    """
    doc = Document(docx_path)
    config = AcronymExpanderConfig()
    processor = AcronymExpander(config)

    processor.doc = doc
    processor._scan_existing_expansions()
    processor._scan_acronyms()
    first_uses = processor._identify_first_uses()

    # Group by acronym
    by_acronym: Dict[str, List] = {}
    for occ in processor.occurrences:
        if occ.acronym not in by_acronym:
            by_acronym[occ.acronym] = []
        by_acronym[occ.acronym].append({
            "para": occ.para_index,
            "first_use": occ.is_first_use,
            "expanded": occ.is_already_expanded,
        })

    return {
        "already_defined": list(processor.defined),
        "needs_expansion": [o.acronym for o in first_uses],
        "by_acronym": by_acronym,
        "unknown": [
            a for a in by_acronym
            if a not in config.known_acronyms and a not in config.organization_acronyms
        ],
    }
