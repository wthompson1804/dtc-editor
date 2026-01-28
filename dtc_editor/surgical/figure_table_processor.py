"""
Figure and Table Caption Processor for Surgical Pipeline

Handles:
1. Detection of figures (w:drawing) and tables in DOCX
2. Detection and validation of existing captions
3. Chapter-based numbering (Figure X-Y format)
4. Source attribution placeholders
5. TOF/TOT regeneration
6. In-text reference correction

Reference: dtc_editor/rules/surgical_rules_manifest.yml (figures_tables section)
"""
from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple, Set
from pathlib import Path
import re
import logging

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class ChapterInfo:
    """Information about a chapter (H1 heading)."""
    para_index: int
    number: int  # 1, 2, 3... (0 for pre-chapter content)
    title: str
    is_numbered: bool  # Whether chapter already has number in text


@dataclass
class FigureInfo:
    """Information about a figure in the document."""
    para_index: int
    chapter_number: int
    sequence_in_chapter: int
    label: str  # "Figure X-Y"
    has_caption: bool
    caption_para_index: Optional[int]
    caption_text: Optional[str]
    has_source: bool
    context_before: str
    context_after: str


@dataclass
class TableInfo:
    """Information about a table in the document."""
    table_index: int
    para_index_before: int  # Paragraph index just before table
    chapter_number: int
    sequence_in_chapter: int
    label: str  # "Table X-Y"
    has_caption: bool
    caption_para_index: Optional[int]
    caption_text: Optional[str]
    has_source: bool


@dataclass
class InTextReference:
    """A reference to a figure/table in body text."""
    para_index: int
    span_start: int
    span_end: int
    original_text: str  # e.g., "Figure 3" or "Figure 1.2"
    ref_type: str  # "figure" or "table"
    corrected_text: Optional[str]  # e.g., "Figure 2-1"


@dataclass
class FigureTableConfig:
    """Configuration for figure/table processing."""
    # Placeholder text for missing captions
    missing_caption_placeholder: str = "INSERT EXPLANATION OF FIGURE OR TABLE"
    missing_source_placeholder: str = "Source: DTC INSERT NAME Working Group."

    # Formatting
    caption_font_size: int = 11
    caption_font_name: str = "Calibri"
    placeholder_color: Tuple[int, int, int] = (255, 0, 0)  # Red for placeholders

    # Behavior
    fix_existing_captions: bool = True
    add_missing_captions: bool = True
    add_missing_sources: bool = True
    fix_in_text_references: bool = True
    regenerate_tof_tot: bool = True

    # Special chapter names that don't get numbers
    unnumbered_chapters: Set[str] = field(default_factory=lambda: {
        "abstract", "references", "authors", "authors & legal notice",
        "appendix", "annex", "acknowledgments", "glossary"
    })


@dataclass
class FigureTableResult:
    """Result of figure/table processing."""
    figures_found: int
    tables_found: int
    captions_added: int
    captions_corrected: int
    sources_added: int
    references_fixed: int
    tof_updated: bool
    tot_updated: bool
    issues: List[str]
    changes: List[Dict]


# =============================================================================
# Main Processor Class
# =============================================================================

class FigureTableProcessor:
    """
    Processes figures and tables in a DOCX document.

    Workflow:
    1. Scan document for chapter structure
    2. Find all figures and tables
    3. Detect existing captions
    4. Generate/correct captions with Chapter-Sequence numbering
    5. Add source attribution where missing
    6. Update TOF/TOT
    7. Fix in-text references
    """

    def __init__(self, config: FigureTableConfig):
        self.config = config
        self.chapters: List[ChapterInfo] = []
        self.figures: List[FigureInfo] = []
        self.tables: List[TableInfo] = []
        self.references: List[InTextReference] = []
        self.doc: Optional[Document] = None

    def process(self, doc: Document) -> FigureTableResult:
        """
        Process all figures and tables in the document.

        Args:
            doc: python-docx Document object

        Returns:
            FigureTableResult with statistics and changes
        """
        self.doc = doc
        self.chapters = []
        self.figures = []
        self.tables = []
        self.references = []

        issues = []
        changes = []

        # Step 1: Scan chapter structure
        self._scan_chapters()
        logger.info(f"Found {len(self.chapters)} chapters")

        # Step 2: Find all figures
        self._scan_figures()
        logger.info(f"Found {len(self.figures)} figures")

        # Step 3: Find all tables
        self._scan_tables()
        logger.info(f"Found {len(self.tables)} tables")

        # Step 4: Find in-text references
        self._scan_references()
        logger.info(f"Found {len(self.references)} in-text references")

        # Step 5: Process figures (add/correct captions)
        captions_added = 0
        captions_corrected = 0
        sources_added = 0

        # Process in reverse order to avoid index shifting
        for figure in reversed(self.figures):
            result = self._process_figure(figure)
            if result.get("caption_added"):
                captions_added += 1
                changes.append({"type": "caption_added", "label": figure.label})
            if result.get("caption_corrected"):
                captions_corrected += 1
                changes.append({"type": "caption_corrected", "label": figure.label})
            if result.get("source_added"):
                sources_added += 1
                changes.append({"type": "source_added", "label": figure.label})
            if result.get("issue"):
                issues.append(result["issue"])

        # Step 6: Process tables
        for table in reversed(self.tables):
            result = self._process_table(table)
            if result.get("caption_added"):
                captions_added += 1
                changes.append({"type": "caption_added", "label": table.label})
            if result.get("caption_corrected"):
                captions_corrected += 1
                changes.append({"type": "caption_corrected", "label": table.label})
            if result.get("source_added"):
                sources_added += 1
                changes.append({"type": "source_added", "label": table.label})
            if result.get("issue"):
                issues.append(result["issue"])

        # Step 7: Fix in-text references
        references_fixed = 0
        if self.config.fix_in_text_references:
            references_fixed = self._fix_references()
            if references_fixed > 0:
                changes.append({"type": "references_fixed", "count": references_fixed})

        # Step 8: Update TOF/TOT
        tof_updated = False
        tot_updated = False
        if self.config.regenerate_tof_tot:
            tof_updated = self._update_tof()
            tot_updated = self._update_tot()
            if tof_updated:
                changes.append({"type": "tof_updated"})
            if tot_updated:
                changes.append({"type": "tot_updated"})

        return FigureTableResult(
            figures_found=len(self.figures),
            tables_found=len(self.tables),
            captions_added=captions_added,
            captions_corrected=captions_corrected,
            sources_added=sources_added,
            references_fixed=references_fixed,
            tof_updated=tof_updated,
            tot_updated=tot_updated,
            issues=issues,
            changes=changes,
        )

    # =========================================================================
    # Scanning Methods
    # =========================================================================

    def _scan_chapters(self) -> None:
        """Scan document for chapter structure (H1 headings)."""
        chapter_num = 0

        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name.lower() if para.style else ""

            # Check for H1 (Heading 1)
            if "heading 1" in style_name or style_name == "heading 1":
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a numbered chapter
                is_numbered = bool(re.match(r'^\d+[\s\.]', text))

                # Check if it's an unnumbered special chapter
                text_lower = text.lower()
                is_special = any(
                    special in text_lower
                    for special in self.config.unnumbered_chapters
                )

                if not is_special:
                    chapter_num += 1

                self.chapters.append(ChapterInfo(
                    para_index=i,
                    number=chapter_num if not is_special else 0,
                    title=text,
                    is_numbered=is_numbered,
                ))

        # If no chapters found, treat whole document as chapter 1
        if not self.chapters:
            self.chapters.append(ChapterInfo(
                para_index=0,
                number=1,
                title="",
                is_numbered=False,
            ))

    def _get_chapter_for_index(self, para_index: int) -> int:
        """Get chapter number for a paragraph index."""
        chapter_num = 1  # Default to chapter 1

        for chapter in self.chapters:
            if chapter.para_index <= para_index and chapter.number > 0:
                chapter_num = chapter.number

        return chapter_num

    def _scan_figures(self) -> None:
        """Scan document for figures (drawings)."""
        # Track figure count per chapter
        chapter_fig_count: Dict[int, int] = {}

        for i, para in enumerate(self.doc.paragraphs):
            # Check for drawings
            drawings = para._element.findall('.//' + qn('w:drawing'))
            if not drawings:
                continue

            # Get chapter number
            chapter_num = self._get_chapter_for_index(i)

            # Increment figure count for this chapter
            if chapter_num not in chapter_fig_count:
                chapter_fig_count[chapter_num] = 0
            chapter_fig_count[chapter_num] += 1
            seq = chapter_fig_count[chapter_num]

            # Generate label
            label = f"Figure {chapter_num}-{seq}"

            # Check for existing caption (next paragraph)
            has_caption = False
            caption_para_index = None
            caption_text = None
            has_source = False

            if i + 1 < len(self.doc.paragraphs):
                next_para = self.doc.paragraphs[i + 1]
                next_text = next_para.text.strip()
                next_style = next_para.style.name.lower() if next_para.style else ""

                # Check if it looks like a caption:
                # 1. Starts with "Figure" followed by number or colon
                # 2. Has "Caption" style
                is_caption_by_text = bool(re.match(r'^Figure\s*[\d\-–:.]', next_text, re.IGNORECASE))
                is_caption_by_style = "caption" in next_style

                if is_caption_by_text or is_caption_by_style:
                    has_caption = True
                    caption_para_index = i + 1
                    caption_text = next_text
                    has_source = "source:" in next_text.lower()

            # Get context
            context_before = ""
            context_after = ""

            for j in range(max(0, i - 3), i):
                txt = self.doc.paragraphs[j].text.strip()
                if txt and not txt.startswith("Figure"):
                    context_before = txt
                    break

            for j in range(i + 1, min(len(self.doc.paragraphs), i + 4)):
                txt = self.doc.paragraphs[j].text.strip()
                if txt and not txt.startswith("Figure"):
                    context_after = txt
                    break

            self.figures.append(FigureInfo(
                para_index=i,
                chapter_number=chapter_num,
                sequence_in_chapter=seq,
                label=label,
                has_caption=has_caption,
                caption_para_index=caption_para_index,
                caption_text=caption_text,
                has_source=has_source,
                context_before=context_before,
                context_after=context_after,
            ))

    def _scan_tables(self) -> None:
        """Scan document for tables."""
        # Track table count per chapter
        chapter_table_count: Dict[int, int] = {}

        # Find paragraph index for each table
        for table_idx, table in enumerate(self.doc.tables):
            # Find the paragraph just before this table
            # This is tricky - we need to find where in the document flow the table appears

            # Get the table's XML element and find its position
            table_elem = table._tbl
            parent = table_elem.getparent()

            # Find paragraph index by looking at document body order
            para_index_before = 0
            for i, para in enumerate(self.doc.paragraphs):
                # Check if this paragraph comes before the table in document order
                para_elem = para._element
                if self._element_comes_before(para_elem, table_elem):
                    para_index_before = i

            # Get chapter number
            chapter_num = self._get_chapter_for_index(para_index_before)

            # Increment table count for this chapter
            if chapter_num not in chapter_table_count:
                chapter_table_count[chapter_num] = 0
            chapter_table_count[chapter_num] += 1
            seq = chapter_table_count[chapter_num]

            # Generate label
            label = f"Table {chapter_num}-{seq}"

            # Check for existing caption (paragraph before table)
            has_caption = False
            caption_para_index = None
            caption_text = None
            has_source = False

            if para_index_before > 0:
                prev_para = self.doc.paragraphs[para_index_before]
                prev_text = prev_para.text.strip()

                if re.match(r'^Table\s+\d', prev_text, re.IGNORECASE):
                    has_caption = True
                    caption_para_index = para_index_before
                    caption_text = prev_text
                    has_source = "source:" in prev_text.lower()

            self.tables.append(TableInfo(
                table_index=table_idx,
                para_index_before=para_index_before,
                chapter_number=chapter_num,
                sequence_in_chapter=seq,
                label=label,
                has_caption=has_caption,
                caption_para_index=caption_para_index,
                caption_text=caption_text,
                has_source=has_source,
            ))

    def _element_comes_before(self, elem1, elem2) -> bool:
        """Check if elem1 comes before elem2 in document order."""
        # Get the body element
        body = self.doc._body._body

        # Find positions
        pos1 = -1
        pos2 = -1

        for idx, child in enumerate(body):
            if child is elem1 or elem1 in child.iter():
                pos1 = idx
            if child is elem2 or elem2 in child.iter():
                pos2 = idx

        return pos1 < pos2 if pos1 >= 0 and pos2 >= 0 else False

    def _scan_references(self) -> None:
        """Scan document for in-text references to figures/tables."""
        # Patterns for references
        fig_pattern = re.compile(
            r'\b(Figure|Fig\.?)\s+(\d+(?:[-–.]\d+)?)',
            re.IGNORECASE
        )
        table_pattern = re.compile(
            r'\b(Table)\s+(\d+(?:[-–.]\d+)?)',
            re.IGNORECASE
        )

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text

            # Skip caption paragraphs
            if re.match(r'^(Figure|Table)\s+\d', text, re.IGNORECASE):
                continue

            # Find figure references
            for match in fig_pattern.finditer(text):
                self.references.append(InTextReference(
                    para_index=i,
                    span_start=match.start(),
                    span_end=match.end(),
                    original_text=match.group(0),
                    ref_type="figure",
                    corrected_text=None,
                ))

            # Find table references
            for match in table_pattern.finditer(text):
                self.references.append(InTextReference(
                    para_index=i,
                    span_start=match.start(),
                    span_end=match.end(),
                    original_text=match.group(0),
                    ref_type="table",
                    corrected_text=None,
                ))

    # =========================================================================
    # Processing Methods
    # =========================================================================

    def _process_figure(self, figure: FigureInfo) -> Dict:
        """Process a single figure - add/correct caption and source."""
        result = {}

        if not figure.has_caption and self.config.add_missing_captions:
            # Add new caption
            self._add_figure_caption(figure)
            result["caption_added"] = True
            result["source_added"] = True  # New captions include source

        elif figure.has_caption:
            # Check and correct existing caption
            if self.config.fix_existing_captions:
                corrected = self._correct_caption(
                    figure.caption_para_index,
                    figure.label,
                    figure.caption_text,
                    "figure"
                )
                if corrected:
                    result["caption_corrected"] = True

            # Add source if missing
            if not figure.has_source and self.config.add_missing_sources:
                self._add_source_to_caption(figure.caption_para_index)
                result["source_added"] = True

        return result

    def _process_table(self, table: TableInfo) -> Dict:
        """Process a single table - add/correct caption and source."""
        result = {}

        if not table.has_caption and self.config.add_missing_captions:
            # Add new caption (tables get caption ABOVE)
            self._add_table_caption(table)
            result["caption_added"] = True
            result["source_added"] = True

        elif table.has_caption:
            # Check and correct existing caption
            if self.config.fix_existing_captions:
                corrected = self._correct_caption(
                    table.caption_para_index,
                    table.label,
                    table.caption_text,
                    "table"
                )
                if corrected:
                    result["caption_corrected"] = True

            # Add source if missing
            if not table.has_source and self.config.add_missing_sources:
                self._add_source_to_caption(table.caption_para_index)
                result["source_added"] = True

        return result

    def _add_figure_caption(self, figure: FigureInfo) -> None:
        """Add a caption paragraph after a figure."""
        # Get the figure paragraph
        fig_para = self.doc.paragraphs[figure.para_index]

        # Create full caption text
        caption_text = (
            f"{figure.label}: {self.config.missing_caption_placeholder}. "
            f"{self.config.missing_source_placeholder}"
        )

        # Insert new paragraph after figure
        new_para = self._insert_paragraph_after(fig_para, caption_text)

        # Style the caption
        self._style_caption_paragraph(new_para, is_placeholder=True)

    def _add_table_caption(self, table: TableInfo) -> None:
        """Add a caption paragraph before a table."""
        # Tables are trickier - we need to insert before the table
        # For now, we'll add after the paragraph before the table

        if table.para_index_before >= 0:
            prev_para = self.doc.paragraphs[table.para_index_before]

            caption_text = (
                f"{table.label}: {self.config.missing_caption_placeholder}. "
                f"{self.config.missing_source_placeholder}"
            )

            new_para = self._insert_paragraph_after(prev_para, caption_text)
            self._style_caption_paragraph(new_para, is_placeholder=True)

    def _correct_caption(
        self,
        para_index: int,
        correct_label: str,
        current_text: str,
        item_type: str,
    ) -> bool:
        """Correct an existing caption to use proper format."""
        if para_index is None or current_text is None:
            return False

        para = self.doc.paragraphs[para_index]

        # Check if format is already correct
        expected_prefix = f"{correct_label}:"
        if current_text.startswith(expected_prefix):
            return False

        # Extract the description part from current caption
        # Match patterns like "Figure 3:", "Figure 1.2:", "Figure 1-2:"
        match = re.match(
            rf'^{item_type}\s*[\d.–-]+\s*:?\s*(.*)$',
            current_text,
            re.IGNORECASE
        )

        if match:
            description = match.group(1).strip()
        else:
            # Couldn't parse - keep original after the label
            description = current_text

        # Build corrected caption
        if description:
            # Ensure description ends with period before source
            if not description.rstrip().endswith('.'):
                description = description.rstrip() + '.'
            new_text = f"{correct_label}: {description}"
        else:
            new_text = f"{correct_label}: {self.config.missing_caption_placeholder}."

        # Update paragraph text
        para.clear()
        run = para.add_run(new_text)

        # Re-apply caption styling
        self._style_caption_paragraph(para, is_placeholder=(self.config.missing_caption_placeholder in new_text))

        return True

    def _add_source_to_caption(self, para_index: int) -> None:
        """Add source placeholder to an existing caption."""
        if para_index is None:
            return

        para = self.doc.paragraphs[para_index]
        current_text = para.text.strip()

        # Ensure text ends with period
        if not current_text.endswith('.'):
            current_text += '.'

        # Add source
        new_text = f"{current_text} {self.config.missing_source_placeholder}"

        # Update paragraph
        para.clear()
        run = para.add_run(new_text)

        # Style the source part in red
        # Note: This is simplified - ideally we'd only color the source part
        self._style_caption_paragraph(para, is_placeholder=True)

    def _insert_paragraph_after(self, para, text: str):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement('w:p')
        para._element.addnext(new_p)

        # Create a new paragraph object
        new_para = self.doc.add_paragraph()
        new_para._element.getparent().remove(new_para._element)

        # Replace the empty element with our new one
        new_p.getparent().replace(new_p, new_para._element)

        # Add the text
        new_para.add_run(text)

        return new_para

    def _style_caption_paragraph(self, para, is_placeholder: bool = False) -> None:
        """Apply caption styling to a paragraph."""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in para.runs:
            run.font.size = Pt(self.config.caption_font_size)
            run.font.name = self.config.caption_font_name

            if is_placeholder:
                run.font.color.rgb = RGBColor(*self.config.placeholder_color)
                run.font.bold = True

    def _fix_references(self) -> int:
        """Fix in-text references to match new figure/table numbers."""
        # Build mapping from old references to new labels
        # This is complex because we need to figure out which figure/table
        # each reference refers to

        # For now, we'll flag references that don't match any known figure/table
        fixed_count = 0

        # Group references by paragraph for efficient processing
        refs_by_para: Dict[int, List[InTextReference]] = {}
        for ref in self.references:
            if ref.para_index not in refs_by_para:
                refs_by_para[ref.para_index] = []
            refs_by_para[ref.para_index].append(ref)

        # Process each paragraph with references
        for para_index, refs in refs_by_para.items():
            para = self.doc.paragraphs[para_index]
            text = para.text

            # Sort refs by position (descending) to avoid offset issues
            refs.sort(key=lambda r: r.span_start, reverse=True)

            for ref in refs:
                # Try to match reference to a known figure/table
                if ref.ref_type == "figure":
                    # Check if reference matches any known figure label
                    matched = False
                    for fig in self.figures:
                        if fig.label.lower() == ref.original_text.lower():
                            matched = True
                            break

                    if not matched:
                        # Reference doesn't match - this could be an old reference
                        # For now, we'll leave it but log it
                        logger.warning(f"Unmatched figure reference: {ref.original_text}")

        return fixed_count

    def _update_tof(self) -> bool:
        """Update Table of Figures with current figure labels."""
        # Find TOF in document
        # TOF is typically marked with a special style or heading "Figures"

        # For now, we'll look for paragraphs with "table of figures" style
        # and update them

        # This is a simplified implementation - full TOF regeneration
        # would require understanding Word's field codes

        return False  # Placeholder - full implementation needed

    def _update_tot(self) -> bool:
        """Update Table of Tables with current table labels."""
        # Similar to TOF
        return False  # Placeholder - full implementation needed


# =============================================================================
# Convenience Function
# =============================================================================

def process_figures_and_tables(
    docx_path: str,
    output_path: str,
    config: Optional[FigureTableConfig] = None,
) -> FigureTableResult:
    """
    Process figures and tables in a DOCX file.

    Args:
        docx_path: Path to input DOCX
        output_path: Path to save output DOCX
        config: Optional configuration

    Returns:
        FigureTableResult with statistics
    """
    if config is None:
        config = FigureTableConfig()

    doc = Document(docx_path)
    processor = FigureTableProcessor(config)
    result = processor.process(doc)

    doc.save(output_path)

    return result
