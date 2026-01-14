from __future__ import annotations
from dataclasses import dataclass
from datetime import datetime
from typing import Optional, Dict, Any
import os
import shutil
import subprocess
import tempfile

@dataclass
class RedlineResult:
    backend: str
    status: str
    message: str
    details: Optional[Dict[str, Any]] = None


def _find_libreoffice() -> Optional[str]:
    """Find LibreOffice executable on the system."""
    # Common installation paths
    lo_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/libreoffice",  # Linux
        "/usr/bin/soffice",  # Linux alternative
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",  # Windows x86
    ]

    for path in lo_paths:
        if os.path.exists(path):
            return path

    # Try system PATH
    return shutil.which("soffice") or shutil.which("libreoffice")


def create_redline(original_docx: str, clean_docx: str, redline_docx: str, author: str = "DTC Editorial Engine", prefer_backend: Optional[str] = None) -> RedlineResult:
    # LibreOffice first (cross-platform, real track changes), then fallbacks
    backends = [prefer_backend] if prefer_backend else ["libreoffice", "aspose", "word_com"]
    last_err = None
    for be in backends:
        if not be:
            continue
        be = be.lower().strip()
        try:
            if be == "libreoffice":
                return _libreoffice_compare(original_docx, clean_docx, redline_docx, author)
            if be == "aspose":
                return _aspose_compare(original_docx, clean_docx, redline_docx, author)
            if be == "word_com":
                return _word_com_compare(original_docx, clean_docx, redline_docx, author)
        except Exception as e:
            last_err = f"{type(e).__name__}: {e}"
            continue
    return RedlineResult(backend="none", status="skipped", message=("No compare backend available. " + (f"Last error: {last_err}" if last_err else "")))


def _libreoffice_compare(original_docx: str, clean_docx: str, redline_docx: str, author: str) -> RedlineResult:
    """Generate redline using visual diff approach.

    Creates a document showing changes with colored text:
    - Deleted text: red strikethrough
    - Added text: blue underline

    This is more reliable than LibreOffice UNO on macOS.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        import difflib
    except ImportError as e:
        raise RuntimeError(f"python-docx required for redline: {e}")

    # Load both documents
    orig_doc = Document(original_docx)
    clean_doc = Document(clean_docx)

    # Create output document based on original structure
    out_doc = Document(original_docx)

    # Get paragraphs from both
    orig_paras = [p.text for p in orig_doc.paragraphs]
    clean_paras = [p.text for p in clean_doc.paragraphs]

    # Match paragraphs and show differences
    matcher = difflib.SequenceMatcher(None, orig_paras, clean_paras)

    # Build a map of which original paragraphs changed and how
    para_changes = {}  # orig_idx -> (change_type, new_text)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        elif tag == 'replace':
            # Paragraphs were modified
            for orig_idx, new_idx in zip(range(i1, i2), range(j1, j2)):
                para_changes[orig_idx] = ('replace', clean_paras[new_idx])
        elif tag == 'delete':
            # Paragraphs were deleted
            for orig_idx in range(i1, i2):
                para_changes[orig_idx] = ('delete', None)
        elif tag == 'insert':
            # New paragraphs were added - mark first original para after
            if i1 < len(orig_paras):
                para_changes[i1] = ('insert_before', clean_paras[j1:j2])

    # Apply changes to output document
    for idx, para in enumerate(out_doc.paragraphs):
        if idx not in para_changes:
            continue

        change_type, new_content = para_changes[idx]

        if change_type == 'delete':
            # Strike through and color red
            para.clear()
            run = para.add_run(orig_paras[idx])
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)

        elif change_type == 'replace':
            orig_text = orig_paras[idx]
            new_text = new_content

            # Clear paragraph and add word-level diff
            para.clear()

            # Word-level diff
            orig_words = orig_text.split()
            new_words = new_text.split()
            word_matcher = difflib.SequenceMatcher(None, orig_words, new_words)

            for tag, i1, i2, j1, j2 in word_matcher.get_opcodes():
                if tag == 'equal':
                    run = para.add_run(' '.join(orig_words[i1:i2]) + ' ')
                elif tag == 'replace':
                    # Show deleted words in red strikethrough
                    if i1 < i2:
                        run = para.add_run(' '.join(orig_words[i1:i2]) + ' ')
                        run.font.strike = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    # Show new words in blue
                    if j1 < j2:
                        run = para.add_run(' '.join(new_words[j1:j2]) + ' ')
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.font.underline = True
                elif tag == 'delete':
                    run = para.add_run(' '.join(orig_words[i1:i2]) + ' ')
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif tag == 'insert':
                    run = para.add_run(' '.join(new_words[j1:j2]) + ' ')
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True

    out_doc.save(redline_docx)

    return RedlineResult(
        backend="python-docx",
        status="ok",
        message="Redline created with visual diff (red=deleted, blue=added)."
    )


def _aspose_compare(original_docx: str, clean_docx: str, redline_docx: str, author: str) -> RedlineResult:
    try:
        import aspose.words as aw  # type: ignore
    except Exception as e:
        raise RuntimeError("Aspose.Words not available.") from e
    base = aw.Document(original_docx)
    revised = aw.Document(clean_docx)
    base.compare(revised, author, datetime.utcnow())
    base.save(redline_docx)
    return RedlineResult(backend="aspose", status="ok", message="Redline created via Aspose compare().")

def _word_com_compare(original_docx: str, clean_docx: str, redline_docx: str, author: str) -> RedlineResult:
    if os.name != "nt":
        raise RuntimeError("Word COM requires Windows.")
    try:
        import win32com.client  # type: ignore
    except Exception as e:
        raise RuntimeError("pywin32 not available.") from e
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        base = word.Documents.Open(os.path.abspath(original_docx), ReadOnly=False)
        revised = word.Documents.Open(os.path.abspath(clean_docx), ReadOnly=True)
        compared = base.Compare(Name=revised.FullName, AuthorName=author)
        compared.SaveAs2(os.path.abspath(redline_docx))
        revised.Close(False); base.Close(False); compared.Close(False)
        return RedlineResult(backend="word_com", status="ok", message="Redline created via Word COM Compare.")
    finally:
        try: word.Quit()
        except Exception: pass
